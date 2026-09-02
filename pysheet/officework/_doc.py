# -*- coding: utf-8 -*-
"""officework.doc — docx のエンジン(Rust)+ python-docx 互換層。

`officework.sheet` と同じ論法で docx を扱う。**原本を正として、変えた所だけ
書き戻す** ので、様式・ヘッダー・図形・変更履歴が壊れない。

    from officework import doc

    d = doc.Doc.open("報告書.docx")
    print(d.unsupported)          # 読めなかった物(空なら取りこぼしなし)
    d[3].text = "差し替え"
    d.replace("旧社名", "新社名")
    d.tables[0].cell(1, 2).text = "125000"   # python-docx の口も通る
    d.save("out.docx")

中身は Rust で、`officework._sheet` が組む1つの拡張の中に副モジュールとして
入っている — maturin が wheel に入れられる拡張は1つなので、`officework.sheet`
と `officework.doc` を**同じ .so に同居させる**ためにこうしてある(利用者に
2つ入れさせない)。この階はそれを包む純 Python の互換層で、python-docx の
口(cell / row_cells / clear / iter_inner_content 等)を足す
(台帳: docs/pysheet-gokan.ja.md)。エンジンには手を入れない。
"""

import os as _os

from . import _sheet as _engine
from ._strict import NoStrayAttributes

_doc = _engine.doc


class _Font(str):
    """Run.font の両対応。

    うちの口では font は**書体名の文字列**(`r.font == "MS明朝"`)。
    python-docx は font という物の下に置く(`r.font.name` / `r.font.size` /
    `r.font.bold`)— 読みも書きも。str の子にして両方を通す。
    書き(`r.font.name = "明朝"`)は元の run(手)に効く。
    """

    def __new__(cls, run):
        self = str.__new__(cls, run.font or "")
        self._run = run
        return self

    @property
    def name(self):
        return str(self) or None

    @name.setter
    def name(self, v):
        self._run.font = v

    @property
    def highlight_color(self):
        """**蛍光ペン**(台帳 #9)。色の名前(`"yellow"` …)か `None`。

        docx の `w:highlight` は決まった色の名前しか受けません。好きな色を
        塗りたいときは背景の塗りを使います。
        """
        return self._run.highlight

    @highlight_color.setter
    def highlight_color(self, v):
        self._run.highlight = None if v is None else str(v)

    @property
    def size(self):
        return self._run.size_pt

    @size.setter
    def size(self, v):
        self._run.size_pt = _to_pt(v)

    @property
    def bold(self):
        return self._run.bold

    @bold.setter
    def bold(self, v):
        # None は「言わない」。python-docx と同じ三択
        self._run.bold = None if v is None else bool(v)

    @property
    def italic(self):
        return self._run.italic

    @italic.setter
    def italic(self, v):
        self._run.italic = None if v is None else bool(v)

    @property
    def underline(self):
        return self._run.underline

    @underline.setter
    def underline(self, v):
        self._run.underline = None if v is None else (bool(v) and v != "none")

    @property
    def color(self):
        """字の色。**本家は色の入れ物を返します**(`font.color.rgb = …`)。

        こちらは字(RRGGBB)で持っているので、字としても `.rgb` でも
        使える入れ物を返します。`font.color = "FF0000"` も通ります
        (2026-08-28、連載のサンプルで踏みました)。
        """
        return _Color(self._run.color, self._run)

    @color.setter
    def color(self, v):
        self._run.color = _rgb_moji(v)

    # **模型が持っている書式は上まで。** 下は docx にある書式で、
    # こちらは読みも書きもしません。python-docx と同じ名前で `None`
    # (=何も言わない)を返します。持っていない物を False と答えると
    # 「切ってある」に読めるので、そうはしません(2026-09-01)。
    @property
    def strike(self):
        return self._run.strike

    @strike.setter
    def strike(self, v):
        self._run.strike = None if v is None else bool(v)

    @property
    def subscript(self):
        return self._run.subscript or None

    @property
    def superscript(self):
        return self._run.superscript or None

    all_caps = property(lambda self: None)
    complex_script = property(lambda self: None)
    cs_bold = property(lambda self: None)
    cs_italic = property(lambda self: None)
    double_strike = property(lambda self: None)
    emboss = property(lambda self: None)
    hidden = property(lambda self: None)
    imprint = property(lambda self: None)
    math = property(lambda self: None)
    no_proof = property(lambda self: None)
    outline = property(lambda self: None)
    rtl = property(lambda self: None)
    shadow = property(lambda self: None)
    small_caps = property(lambda self: None)
    snap_to_grid = property(lambda self: None)
    spec_vanish = property(lambda self: None)
    web_hidden = property(lambda self: None)


def _rgb_moji(v):
    """色を `RRGGBB` の字にする。本家の `RGBColor(255,0,0)` も受けます"""
    if v is None:
        return None
    if isinstance(v, _Color):
        return str(v) or None
    # RGBColor は 3 バイトの列。str() が "FF0000" を返します
    t = str(v).strip().lstrip("#")
    if len(t) == 8:
        t = t[2:]
    return t or None


def _muki(v):
    """向きの言い方をそろえる。字・本家の列挙・数(0=縦, 1=横)を受けます"""
    if isinstance(v, int) and not isinstance(v, bool):
        return "landscape" if v == 1 else "portrait"
    t = str(getattr(v, "name", v)).strip().lower()
    if t.startswith("landscape"):
        return "landscape"
    if t.startswith("portrait"):
        return "portrait"
    return t


def _kumikomi_kind(name):
    """組み込みスタイルの種類。名前が `Char` で終われば文字スタイル"""
    return "character" if str(name).strip().endswith("Char") else "paragraph"


class _Color(str):
    """字の色。**字としても `.rgb` でも使えます。**

    本家(python-docx)は `run.font.color.rgb = RGBColor(255,0,0)` と
    書きます。こちらは色を字で持っているので、字を継いで `.rgb` を
    足しました。`_Font` の手と同じです。
    """

    def __new__(cls, value, run=None):
        self = super().__new__(cls, value or "")
        self._run = run
        return self


    @property
    def theme_color(self):
        """テーマの色。**模型は RRGGBB で持ちます**ので、常に None です。"""
        return None
    @property
    def rgb(self):
        return str(self) or None

    @rgb.setter
    def rgb(self, v):
        if self._run is None:
            raise NotImplementedError("この色は run に繋がっていません")
        self._run.color = _rgb_moji(v)

    @property
    def type(self):
        """色の指し方。うちは常に直の指定(本家の MSO_THEME_COLOR は無し)"""
        return None


class Run:
    """書式のまとまり。**位置で引き直す手**(python-docx の run と同じ使い方)。
    `r.bold = True` も `r.add_text("続き")` も効く。段落の text の代入や
    replace で run の並びが変わった後は、`runs` から引き直すこと。"""

    __slots__ = ("_r",)

    def __init__(self, raw):
        self._r = raw

    @property
    def text(self):
        return self._r.text

    @text.setter
    def text(self, v):
        self._r.text = v

    @property
    def size_pt(self):
        """字の大きさ(pt)。None = 指定なし(様式・文書の既定に従う)。"""
        return self._r.size_pt

    @size_pt.setter
    def size_pt(self, v):
        # setter だけが包みから漏れていた(2026-08-14 に発見。エンジン側には
        # 2026-08-12 からある)。None で指定を外せるのはエンジンと同じ約束
        self._r.size_pt = None if v is None else float(v)

    @property
    def font(self):
        return _Font(self._r)

    @property
    def bold(self):
        return self._r.bold

    @bold.setter
    def bold(self, v):
        self._r.bold = None if v is None else bool(v)

    @property
    def italic(self):
        return self._r.italic

    @italic.setter
    def italic(self, v):
        self._r.italic = None if v is None else bool(v)

    @property
    def underline(self):
        return self._r.underline

    @underline.setter
    def underline(self, v):
        self._r.underline = None if v is None else (bool(v) and v != "none")

    @property
    def strike(self):
        return self._r.strike

    @strike.setter
    def strike(self, v):
        self._r.strike = None if v is None else bool(v)

    @property
    def contains_page_break(self):
        """この run に改ページが入っているか。

        **模型は改ページを段落の性質で持ちます**。run では持たないので、
        常に False です(段落の側は `p.contains_page_break` で見られます)。
        """
        return False

    def add_picture(self, image, width=None, height=None):
        """この run の段落に画像を足す(python-docx と同じ口)。

        大きさは mm の数でも、本家の `Mm(15)` のような Length でも。
        """
        self._r.add_picture(image, _to_mm(width) or None, _to_mm(height) or None)

    @property
    def highlight(self):
        """蛍光ペンの色の名前。`run.font.highlight_color` の短い書き方"""
        return self._r.highlight

    @highlight.setter
    def highlight(self, v):
        self._r.highlight = None if v in (None, "none") else str(v)

    @property
    def superscript(self):
        """上付き(x²)"""
        return self._r.superscript

    @superscript.setter
    def superscript(self, v):
        self._r.superscript = bool(v)

    @property
    def subscript(self):
        """下付き(H₂O)"""
        return self._r.subscript

    @subscript.setter
    def subscript(self, v):
        self._r.subscript = bool(v)

    @property
    def color(self):
        return self._r.color

    @color.setter
    def color(self, v):
        self._r.color = v

    @property
    def style(self):
        """文字スタイルの名前(指定なしは None)。字としても `.name` でも
        読めます。書きは styles にある文字スタイルの名前(本家の物でも)。"""
        n = self._r.style
        # 本家は指定が無いと "Default Paragraph Font" を返します
        return StyleName(n or "Default Paragraph Font", kind="character")

    @style.setter
    def style(self, v):
        self._r.style = None if v is None else str(getattr(v, "name", v))

    @property
    def hyperlink(self):
        """リンク先(URL。無ければ None)。"""
        return self._r.hyperlink

    @hyperlink.setter
    def hyperlink(self, v):
        self._r.hyperlink = v

    def add_break(self, break_type=None):
        """改行を足す(python-docx と同じ口)。docx の w:br になる。
        改ページの break_type は段落の性質(page_break_before)で持つので断る。"""
        if break_type is not None:
            raise NotImplementedError(
                "改ページは段落の性質(paragraph_format.page_break_before)で持つ(台帳)"
            )
        self._r.add_break()

    def add_tab(self):
        """タブを足す(python-docx と同じ口)。docx の w:tab になる。"""
        self._r.add_tab()

    def iter_inner_content(self):
        """run の中身を順に(本家と同じ口)。字は str、改行とタブは
        Break / Tab で返す — うちは両方を run の字(\\n・\\t)で持つので、
        ここで**順のまま**解いて見せる。"""
        buf = ""
        for ch in self._r.text:
            if ch in "\n\t":
                if buf:
                    yield buf
                    buf = ""
                yield Break() if ch == "\n" else Tab()
            else:
                buf += ch
        if buf:
            yield buf

    def mark_comment_range(self, last_run, comment_id):
        """本家は run から run までを範囲としてコメントに紐づけるが、
        **うちのコメントは段落単位**(模型の粒度)。範囲は持てないので断る —
        段落に付けるなら Paragraph.add_comment。"""
        raise NotImplementedError(
            "コメントの範囲は段落単位(模型の粒度)。Paragraph.add_comment を使う(台帳)"
        )

    def add_text(self, text):
        """字を後ろに継ぎ足す(書式はこの run のまま — 本家と同じ定義)。"""
        self._r.add_text(text)

    def clear(self):
        """字を消す(書式は残る)。返りは自分(本家と同じ)。"""
        self._r.clear()
        return self

    def __repr__(self):
        return repr(self._r)


class Break:
    """run の中の改行(docx の w:br)。iter_inner_content が返す。"""

    __slots__ = ()

    def __repr__(self):
        return "<officework.doc Break>"


class Tab:
    """run の中のタブ(docx の w:tab)。"""

    __slots__ = ()

    def __repr__(self):
        return "<officework.doc Tab>"


class Length(int):
    """docx の長さ(EMU)。本家の Length と同じ算術(.mm / .cm / .pt / .emu)。"""

    @property
    def emu(self):
        return int(self)

    @property
    def mm(self):
        return self / 36000

    @property
    def cm(self):
        return self / 360000

    @property
    def pt(self):
        return self / 12700

    @property
    def inches(self):
        return self / 914400

    @property
    def twips(self):
        return self / 635

    @classmethod
    def from_inches(cls, v):
        return None if v is None else cls(round(float(v) * 914400))

    @classmethod
    def from_pt(cls, v):
        """pt → Length。`None` はそのまま `None`(指定なし)。"""
        return None if v is None else cls(round(float(v) * 12700))

    @classmethod
    def from_mm(cls, v):
        return None if v is None else cls(round(float(v) * 36000))


def _to_pt(v):
    """`Length` でも生の数でも pt にする。**本家は Length、うちは pt** —
    どちらで渡されても通します(`Pt(12)` も `12` も同じ意味)。

    **型では見ません。** 本家(python-docx)の `Pt(12)` はあちらの
    `Length` で、こちらの `Length` とは別の型です。中身は EMU の整数
    なので、型で見ると `152400` がそのまま pt として渡り「大きさが変」で
    止まります(2026-08-28、連載のサンプルで踏みました)。`.pt` を
    持っているかどうかで見ます。
    """
    if v is None:
        return 0.0
    pt = getattr(v, "pt", None)
    return float(pt) if pt is not None else float(v)


def _to_mm(v):
    if v is None:
        return 0.0
    mm = getattr(v, "mm", None)
    return float(mm) if mm is not None else float(v)


def Pt(v):
    """pt → Length。本家の docx.shared.Pt と同じ。"""
    return Length(round(v * 12700))


def Mm(v):
    """mm → Length。本家の docx.shared.Mm と同じ。"""
    return Length(round(v * 36000))


def Cm(v):
    """cm → Length。本家の docx.shared.Cm と同じ。"""
    return Length(round(v * 360000))


def Inches(v):
    """インチ → Length。本家の docx.shared.Inches と同じ。"""
    return Length(round(v * 914400))


def Emu(v):
    """EMU をそのまま Length に。本家の docx.shared.Emu と同じ。"""
    return Length(int(v))


def Twips(v):
    """twip → Length。本家の docx.shared.Twips と同じ(1pt = 20twip)。"""
    return Length(round(v * 635))


# **python-docx が入っていれば、その `RGBColor` を土台にします**
# (2026-09-01 発注者)。本家の口は `isinstance` で自分の型かどうかを
# 見るので、別の型だと「RGBColor object でない」と断られます。移り
# 変わる途中では、本家で組みながら色だけこちらから取る書き方をします。
# 入っていなければ、こちらだけで同じ物を作ります
try:  # pragma: no cover - python-docx の有無で分かれます
    from docx.shared import RGBColor as _hon_rgb
    _hon_aru = True
except Exception:  # pragma: no cover
    _hon_rgb = tuple
    _hon_aru = False


class RGBColor(_hon_rgb):
    """字の色。本家の `docx.shared.RGBColor` と同じ使い方です。

    `RGBColor(0xFF, 0x00, 0x00)` で作り、`str()` は `"FF0000"` です。
    こちらは色を `RRGGBB` の字で持つので、そのまま渡せます。
    """

    __slots__ = ()

    def __new__(cls, r, g, b):
        for v in (r, g, b):
            if not isinstance(v, int) or not 0 <= v <= 255:
                raise ValueError("RGBColor の各成分は 0〜255 の整数です")
        if _hon_aru:
            return super().__new__(cls, r, g, b)
        return super().__new__(cls, (r, g, b))

    @classmethod
    def from_string(cls, s):
        """`"FF0000"` から作ります(本家と同じ)。"""
        t = str(s).strip().lstrip("#")
        if len(t) != 6:
            raise ValueError("色は RRGGBB の6桁です: {!r}".format(s))
        return cls(int(t[0:2], 16), int(t[2:4], 16), int(t[4:6], 16))

    def __str__(self):
        return "{:02X}{:02X}{:02X}".format(*self)

    def __repr__(self):
        return "RGBColor(0x{:02x}, 0x{:02x}, 0x{:02x})".format(*self)


def Document(path=None, lang=None):
    """空の文書を作る、または開く。**本家の `docx.Document` と同じ名前**です。

        from officework import doc
        d = doc.Document()            # 空の文書
        d = doc.Document("報告.docx")  # 開く

    中身は [`Doc`] と同じです。本家の見本をそのまま持ってくるための
    別名です(2026-09-01。見本が `docx.Document()` で書いてありました)。
    """
    return Doc(path, lang)


class InlineShape:
    """文書の中の画像(本家の InlineShape の役)。width / height は Length。"""

    __slots__ = ("width", "height")

    def __init__(self, w_mm, h_mm):
        self.width = Mm(w_mm)
        self.height = Mm(h_mm)

    def __repr__(self):
        return "<officework.doc InlineShape {:.0f}×{:.0f}mm>".format(
            self.width.mm, self.height.mm)


class Hyperlink:
    """段落の中のリンク(本家の Hyperlink の役)。text と address。"""

    __slots__ = ("text", "address")

    def __init__(self, text, address):
        self.text = text
        self.address = address

    def __repr__(self):
        return "<officework.doc Hyperlink {!r} → {}>".format(self.text, self.address)


class Section:
    """節(本家の Section の役)。紙の大きさ・余白は Length(EMU)で
    読み書き — 書きは原文の sectPr へ属性差し替えなので、理解しない設定
    (ヘッダー参照・段組み)は崩れない。"""

    __slots__ = ("_s", "_doc")

    def __init__(self, raw, doc=None):
        self._s = raw
        self._doc = doc

    def _len_prop(name):  # noqa: N805 — 小さな工場
        mm_name = name + "_mm"

        def get(self):
            return Mm(getattr(self._s, mm_name))

        def set_(self, v):
            setattr(self._s, mm_name, v.mm if hasattr(v, "mm") else float(v) / 36000)

        return property(get, set_)

    @property
    def start_type(self):
        """節の始め方。"new_page" か "continuous"(python-docx と同じ)"""
        return self._s.start_type

    @property
    def header(self):
        """この節のヘッダー。

        **模型はヘッダーを文書に1つ持ちます**(節ごとではありません)。
        docx の途中の節が別のヘッダーを持つ形は、原文のまま持ち越して
        いて、こちらからは触りません。ここが返すのは文書のヘッダーです。
        """
        return self._doc.header if self._doc is not None else None

    @property
    def footer(self):
        """この節のフッター。ヘッダーと同じく文書の物を返します"""
        return self._doc.footer if self._doc is not None else None

    page_width = _len_prop("page_width")
    page_height = _len_prop("page_height")
    left_margin = _len_prop("left_margin")
    right_margin = _len_prop("right_margin")
    top_margin = _len_prop("top_margin")
    bottom_margin = _len_prop("bottom_margin")
    del _len_prop

    @property
    def orientation(self):
        return self._s.orientation

    @orientation.setter
    def orientation(self, v):
        """`"portrait"` / `"landscape"`。**幅と高さを1手で入れ替えます** —
        1つずつ動かすと、途中で正方形になって向きが決まらない瞬間ができます。

        本家の列挙(`WD_ORIENT.LANDSCAPE`)でも受けます。`str()` すると
        `"LANDSCAPE (1)"` になるので、名前と数の方を先に見ます
        (2026-08-28、連載のサンプルで踏みました)。
        """
        self._s.orientation = _muki(v)

    def __repr__(self):
        return repr(self._s)


class _StylePara:
    """`Style.paragraph_format` の役。段落の `ParagraphFormat` と同じ呼び名。

    **スタイルが持つ段落の見た目**です(docx の `w:pPr`)。段落そのものの
    書式とは別で、こちらはスタイルを当てた段落すべてに効きます。
    """

    __slots__ = ("_s",)

    def __init__(self, style):
        self._s = style

    @property
    def alignment(self):
        return self._s._props()["alignment"]

    @alignment.setter
    def alignment(self, v):
        self._s._set(alignment=None if v is None else _align_word(v))

    @property
    def space_before(self):
        v = self._s._props()["space_before"]
        return None if v is None else Length.from_pt(v)

    @space_before.setter
    def space_before(self, v):
        self._s._set(space_before=None if v is None else _to_pt(v))

    @property
    def space_after(self):
        v = self._s._props()["space_after"]
        return None if v is None else Length.from_pt(v)

    @space_after.setter
    def space_after(self, v):
        self._s._set(space_after=None if v is None else _to_pt(v))

    @property
    def line_spacing(self):
        return self._s._props()["line_spacing"]

    @line_spacing.setter
    def line_spacing(self, v):
        self._s._set(line_spacing=None if v is None else float(v))

    @property
    def left_indent(self):
        """左の字下げ。**模型は段数**(1段=全角2字)なので pt に直します"""
        v = self._s._props()["indent_level"]
        return None if v is None else Length.from_pt(v * _ZEN * 2)

    @left_indent.setter
    def left_indent(self, v):
        if v is None:
            self._s._set(indent_level=None)
            return
        self._s._set(indent_level=max(0, min(9, round(_to_pt(v) / (_ZEN * 2)))))

    @property
    def first_line_indent(self):
        v = self._s._props()["first_line_indent"]
        return None if v is None else Length.from_pt(v)

    @first_line_indent.setter
    def first_line_indent(self, v):
        self._s._set(first_line_indent=None if v is None else _to_pt(v))


class _StyleColor(str):
    """スタイルの字の色。字としても `.rgb` でも使えます(`_Color` と同じ手)"""

    def __new__(cls, value, font=None):
        self = super().__new__(cls, value or "")
        self._font = font
        return self


    @property
    def theme_color(self):
        """テーマの色。**模型は RRGGBB で持ちます**ので、常に None です。"""
        return None
    @property
    def rgb(self):
        return str(self) or None

    @rgb.setter
    def rgb(self, v):
        if self._font is None:
            raise NotImplementedError("この色はスタイルに繋がっていません")
        self._font._set(color=_rgb_moji(v))

    @property
    def type(self):
        return None


class _StyleFont:
    """`Style.font` の役。**自作スタイルだけ**書けます(2026-08-27)。

    原本から読んだスタイルの定義は据え置きで持ち越すので、触ると原本の
    様式が崩れます。書こうとすると断って理由を言います。
    """

    __slots__ = ("_d", "_name")

    def __init__(self, raw_doc, name):
        self._d = raw_doc
        self._name = name

    def _look(self):
        if self._d is None:
            return {}
        try:
            return self._d.style_look(self._name) or {}
        except Exception:
            return {}

    def _set(self, **kw):
        self._d.set_style_look(self._name, **kw)

    @property
    def bold(self):
        return self._look().get("bold")

    @bold.setter
    def bold(self, v):
        self._set(bold=None if v is None else bool(v))

    @property
    def italic(self):
        return self._look().get("italic")

    @italic.setter
    def italic(self, v):
        self._set(italic=None if v is None else bool(v))

    @property
    def underline(self):
        return self._look().get("underline")

    @underline.setter
    def underline(self, v):
        self._set(underline=None if v is None else bool(v))

    @property
    def strike(self):
        return self._look().get("strike")

    @strike.setter
    def strike(self, v):
        self._set(strike=None if v is None else bool(v))

    @property
    def size(self):
        """字の大きさ。**Length で返します**(`.pt` で pt になります)。"""
        return Length.from_pt(self._look().get("size"))

    @size.setter
    def size(self, v):
        self._set(size=None if v is None else _to_pt(v))

    @property
    def color(self):
        """字の色。**本家は色の入れ物を返します**(`font.color.rgb`)。

        色を持たないスタイルでも入れ物を返します — `None` を返すと
        `font.color.rgb` が読めません(2026-08-28、連載の第3回)。
        """
        return _StyleColor(self._look().get("color"), self)

    @color.setter
    def color(self, v):
        self._set(color=_rgb_moji(v))

    @property
    def name(self):
        """書体の名前(スタイルの名前ではありません — 本家と同じ)。"""
        return self._look().get("font")

    @name.setter
    def name(self, v):
        self._set(font=None if v is None else str(v))

    # **模型が持っていない書式。** python-docx と同じ名前で None を返します
    all_caps = property(lambda self: None)
    complex_script = property(lambda self: None)
    cs_bold = property(lambda self: None)
    cs_italic = property(lambda self: None)
    double_strike = property(lambda self: None)
    emboss = property(lambda self: None)
    hidden = property(lambda self: None)
    highlight_color = property(lambda self: None)
    imprint = property(lambda self: None)
    math = property(lambda self: None)
    no_proof = property(lambda self: None)
    outline = property(lambda self: None)
    rtl = property(lambda self: None)
    shadow = property(lambda self: None)
    small_caps = property(lambda self: None)
    snap_to_grid = property(lambda self: None)
    spec_vanish = property(lambda self: None)
    subscript = property(lambda self: None)
    superscript = property(lambda self: None)
    web_hidden = property(lambda self: None)

    def __repr__(self):
        return "<officework.doc StyleFont {!r}>".format(self._look())


def _style_na(doc, sid):
    """styleId から、表に出る名前へ。引けなければ id のまま。"""
    try:
        return doc.style_props(sid).get("name") or sid
    except Exception:
        return sid


class _TabStops:
    """タブの止まる位置の一覧。**模型が持っていない**ので、いつも空です。"""

    __slots__ = ()

    def __iter__(self):
        return iter(())

    def __len__(self):
        return 0

    def __repr__(self):
        return "<officework.doc TabStops []>"


class _StyleParagraphFormat:
    """`Style.paragraph_format` の役。**読むだけ**です。
    スタイルが言っていない所は None を返します(0 ではありません)。"""

    __slots__ = ("_d",)

    def __init__(self, d):
        self._d = d or {}

    @property
    def alignment(self):
        return self._d.get("alignment")

    @property
    def first_line_indent(self):
        v = self._d.get("first_line_indent")
        return None if v is None else Length.from_pt(v)

    @property
    def left_indent(self):
        n = self._d.get("indent_level")
        return None if not n else Length.from_pt(n * _ZEN * 2)

    @property
    def space_before(self):
        v = self._d.get("space_before")
        return None if v is None else Length.from_pt(v)

    @property
    def space_after(self):
        v = self._d.get("space_after")
        return None if v is None else Length.from_pt(v)

    @property
    def line_spacing(self):
        return self._d.get("line_spacing")

    # 模型が持っていない性質
    keep_together = property(lambda self: None)
    keep_with_next = property(lambda self: None)
    line_spacing_rule = property(lambda self: None)
    page_break_before = property(lambda self: None)
    right_indent = property(lambda self: None)
    widow_control = property(lambda self: None)
    @property
    def tab_stops(self):
        """タブを打ったとき字が止まる位置。**模型は持っていません**ので、
        いつも空です。足す口も今はありません。"""
        return _TabStops()


class StyleName(str):
    """スタイルの名前。**字としても Style としても振る舞います。**

    本家(python-docx)の `paragraph.style` は Style を返すので
    `p.style.name` と書きます。こちらは名前(字)を返していたので、
    その書き方が全部止まりました(2026-08-28、連載のサンプル)。

    `p.style == "Title"` と `p.style.name == "Title"` の**どちらも**
    通るように、字を継いで `.name` を足しました。`_Font` が
    `run.font` で使っているのと同じ手です。
    """

    def __new__(cls, name, doc=None, kind="paragraph", sid=None):
        self = super().__new__(cls, name or "")
        self._doc = doc
        self._kind = kind
        self._sid = sid
        return self

    @property
    def name(self):
        return str(self)

    @property
    def style_id(self):
        """docx の中の名前(`w:styleId`)。分からなければ表に出る名前。"""
        return self._sid or str(self)

    @property
    def type(self):
        return self._kind

    def _props(self):
        if self._doc is None:
            return {}
        try:
            return self._doc.style_props(str(self))
        except Exception:
            return {}

    @property
    def font(self):
        """このスタイルが言っている字の見た目。"""
        if self._doc is None:
            return _StyleFont(None, str(self))
        return _StyleFont(self._doc, str(self))

    @property
    def base_style(self):
        """元になるスタイル(docx の `w:basedOn`)。無ければ None。"""
        b = self._props().get("based_on")
        if b is None:
            return None
        # `w:basedOn` は styleId です。表に出る名前に直します
        return StyleName(_style_na(self._doc, b), self._doc, self._kind, sid=b)

    @property
    def builtin(self):
        """Word に元からあるスタイルか。名前で見分けます。"""
        return str(self) in _KUMIKOMI

    @property
    def hidden(self):
        return self._props().get("hidden")

    @property
    def locked(self):
        return self._props().get("locked")

    @property
    def quick_style(self):
        return self._props().get("quick_style")

    @property
    def unhide_when_used(self):
        return self._props().get("unhide_when_used")

    @property
    def priority(self):
        return self._props().get("priority")

    # **模型が持っていない性質。** python-docx と同じ名前で None を返します
    # (持っていない物を False と答えると「切ってある」に読めます)
    @property
    def next_paragraph_style(self):
        """次の段落のスタイル。docx が言っていなければ自分自身(本家と同じ)。"""
        return self
    @property
    def paragraph_format(self):
        """このスタイルが言っている段落の形。言っていない所は None。"""
        return _StyleParagraphFormat(self._props())


_KUMIKOMI = frozenset({
    "Normal", "Body Text", "Title", "Subtitle", "Caption", "Quote",
    "List Paragraph", "List Number", "List Bullet", "Header", "Footer",
    "Default Paragraph Font", "No Spacing", "Table Grid",
} | {f"Heading {i}" for i in range(1, 10)})


class Style:
    """スタイルの名乗り(本家の style の役 — .name / .style_id / .type)と、
    自作スタイルの見た目(`.font`)。

    原本から読んだスタイルの定義は styles.xml が持ち、保存で原本のまま
    持ち越されます。**自作した物だけ** `.font` で書けます(2026-08-27)。"""

    __slots__ = ("style_id", "name", "type", "_d")

    def __init__(self, style_id, name, kind, raw_doc=None):
        self.style_id = style_id
        self.name = name
        self.type = kind
        self._d = raw_doc

    # ── 定義の性質(python-docx と同じ呼び名)────────────────────────
    #
    # `base_style` は本家に合わせて**スタイルそのもの**を返します。
    # 他は真偽と数です。原本から読んだスタイルでも書けて、保存では
    # 触った定義だけが styles.xml で差し替わります(2026-08-28)。

    def _props(self):
        if self._d is None:
            raise NotImplementedError("このスタイルは文書に繋がっていません")
        return self._d.style_props(self.name)

    def _set(self, **kw):
        if self._d is None:
            raise NotImplementedError("このスタイルは文書に繋がっていません")
        self._d.set_style_props(self.name, **kw)

    @property
    def base_style(self):
        b = self._props()["based_on"]
        if b is None:
            return None
        for sid, name, kind in self._d.styles:
            if sid == b:
                return Style(sid, name, kind, self._d)
        return None

    @base_style.setter
    def base_style(self, v):
        if v is None:
            self._set(based_on=None)
            return
        # スタイルでも名前でも受けます
        sid = getattr(v, "style_id", None)
        if sid is None:
            name = str(getattr(v, "name", v))
            sid = next((i for i, n, _ in self._d.styles if n == name or i == name), name)
        self._set(based_on=sid)

    @property
    def hidden(self):
        return self._props()["hidden"]

    @hidden.setter
    def hidden(self, v):
        self._set(hidden=bool(v))

    @property
    def unhide_when_used(self):
        return self._props()["unhide_when_used"]

    @unhide_when_used.setter
    def unhide_when_used(self, v):
        self._set(unhide_when_used=bool(v))

    @property
    def locked(self):
        return self._props()["locked"]

    @locked.setter
    def locked(self, v):
        self._set(locked=bool(v))

    @property
    def quick_style(self):
        return self._props()["quick_style"]

    @quick_style.setter
    def quick_style(self, v):
        self._set(quick_style=bool(v))

    def delete(self):
        """このスタイルを消す(python-docx と同じ口)。

        **このアプリで足した物だけ**です。原本から読んだスタイルは
        据え置きなので、消そうとすると断ります。
        """
        if self._d is None:
            raise NotImplementedError("このスタイルは文書に繋がっていません")
        self._d.remove_style(self.name)

    @property
    def paragraph_format(self):
        """スタイルの段落の見た目(揃え・前後の空き・行間)"""
        return _StylePara(self)

    @property
    def priority(self):
        return self._props()["priority"]

    @priority.setter
    def priority(self, v):
        self._set(priority=None if v is None else int(v))

    @property
    def font(self):
        if self._d is None:
            raise NotImplementedError("このスタイルは文書に繋がっていません")
        return _StyleFont(self._d, self.name)

    def __repr__(self):
        return "<officework.doc Style {!r} ({})>".format(self.name, self.type)


# 模型の様式名 → docx の様式。**模型は本文を "body" と呼ぶ**が、docx では
# "Normal"。ここを繋がないと、段落から読んだ名前で styles[…] が引けない
_STYLE_ALIAS = {"body": "normal", "normal": "body"}


def _style_key(name):
    """様式名の照合の形。docx は style_id("Heading1")と UI 名("heading 1")の
    2つの名乗りを持ち、模型はさらに "heading1" と呼ぶ — **どれで引いても
    同じ様式に当たる**ようにする(本家も id と UI 名の両方で引ける)。"""
    return str(name).replace(" ", "").replace("-", "").lower()


class _Styles:
    """Doc.styles の返り(本家の Styles の役)。名前で引け、add_style で足せる。"""

    def __init__(self, raw_doc):
        self._d = raw_doc

    def _all(self):
        return [Style(i, n, k, self._d) for i, n, k in self._d.styles]

    def __iter__(self):
        return iter(self._all())

    def __len__(self):
        return len(self._all())

    def _find(self, name):
        k = _style_key(name)
        keys = {k, _STYLE_ALIAS.get(k, k)}
        for s in self._all():
            if {_style_key(s.name), _style_key(s.style_id)} & keys:
                return s
        return None

    def __contains__(self, name):
        return self._find(name) is not None

    def __getitem__(self, name):
        s = self._find(name)
        if s is not None:
            return s
        # **Word が使ったときに作る組み込みスタイル**なら、ここで作ります。
        # 段落に貼るときと同じ作法です(2026-08-28)
        try:
            return self.add_style(str(name), _kumikomi_kind(name), builtin=True)
        except Exception:
            pass
        raise KeyError("スタイルが無い: {!r}".format(name))

    def add_style(self, name, style_type="paragraph", builtin=False):
        """スタイルを足す(本家と同じ口)。style_type は "paragraph" /
        "character" / "table"(本家の WD_STYLE_TYPE でもよい)。
        名乗りだけの最小定義 — 見た目は直接書式が第一のまま。"""
        kind = str(getattr(style_type, "name", style_type)).lower()
        self._d.add_style(name, kind)
        return self[name]

    def __repr__(self):
        return "<officework.doc Styles {}>".format([s.name for s in self._all()])


class Comment:
    """段落に付いたコメント(本家の Comment の役)。段落単位の粒度。"""

    __slots__ = ("author", "text", "paragraph")

    def __init__(self, author, text, paragraph):
        self.author = author
        self.text = text
        self.paragraph = paragraph

    def __repr__(self):
        return "<officework.doc Comment {!r} by {!r}>".format(self.text, self.author)


def _align_word(v):
    # 寄せの受け皿: "center" / python-docx の WD_ALIGN_PARAGRAPH(.name)/ None
    if v is None:
        return None
    return str(getattr(v, "name", v)).lower()


# 全角1字の幅(pt)。本文の既定の大きさ 10.5pt で数えます
_ZEN = 10.5


class ParagraphFormat:
    """python-docx の paragraph_format の役。模型が持つ物だけ —
    alignment・line_spacing・page_break_before・space_before・space_after・
    first_line_indent・left_indent。

    **左の字下げは模型では段数**(1段=全角2字)です。python-docx は長さで
    書くので、本文の既定の大きさ(10.5pt)で数えて行き来します。段の途中の
    値はいちばん近い段に寄ります(2026-08-27 に決めました)。1行目の
    字下げ(first_line_indent)は docx と同じ長さで持つので、そのままです。"""

    __slots__ = ("_p",)

    def __init__(self, raw):
        self._p = raw

    @property
    def alignment(self):
        return self._p.align

    @alignment.setter
    def alignment(self, v):
        self._p.align = _align_word(v) or "left"

    @property
    def line_spacing(self):
        return self._p.line_spacing

    @line_spacing.setter
    def line_spacing(self, v):
        self._p.line_spacing = float(v)

    @property
    def line_spacing_rule(self):
        """行間の決め方。"auto" / "exact" / "atLeast"。無指定は None。"""
        return self._p.line_spacing_rule

    @property
    def right_indent(self):
        """右の字下げ。**まだ模型が持っていません**ので常に None です。

        代入は受けますが、捨てます。断ると本家の台本が途中で止まり、
        持っている物まで書けなくなるためです(2026-09-01)。読めなかった
        物は `unsupported` と同じ扱いで、`d.unsupported` に出ます。
        """
        return None

    @right_indent.setter
    def right_indent(self, v):
        self._p.note_unsupported("段落の右の字下げ(right_indent)")

    @property
    def page_break_before(self):
        return self._p.page_break_before

    @page_break_before.setter
    def page_break_before(self, v):
        self._p.page_break_before = bool(v)

    @property
    def first_line_indent(self):
        """1行目の字下げ。正で字下げ、負でぶら下げ(python-docx と同じ)。
        文書が何も言っていなければ None。"""
        v = self._p.first_line_indent
        return None if v is None else Length.from_pt(v)

    @first_line_indent.setter
    def first_line_indent(self, v):
        self._p.first_line_indent = 0.0 if v is None else _to_pt(v)

    @property
    def left_indent(self):
        """左の字下げ。docx の `w:ind w:left` をそのまま返します。
        文書が何も言っていなければ None。

        段数(1段=全角2字)で持つ口は `p.indent_level` です。"""
        v = self._p.left_indent
        if v is None:
            n = self._p.indent_level
            return Length.from_pt(n * _ZEN * 2) if n else None
        return Length.from_pt(v)

    @left_indent.setter
    def left_indent(self, v):
        # twip をそのまま持ちます。段数の口も揃えておきます
        pt = None if v is None else _to_pt(v)
        self._p.left_indent = pt
        self._p.indent_level = 0 if pt is None else max(0, min(9, round(pt / (_ZEN * 2))))

    @property
    def space_before(self):
        """段落の前の空き。**pt で返します**(python-docx は Length)。"""
        v = self._p.space_before
        return None if v is None else Length.from_pt(v)

    @space_before.setter
    def space_before(self, v):
        self._p.space_before = _to_pt(v)

    @property
    def space_after(self):
        v = self._p.space_after
        return None if v is None else Length.from_pt(v)

    @space_after.setter
    def space_after(self, v):
        self._p.space_after = _to_pt(v)


class Paragraph:
    """1つの段落。本文にも表のセルの中にもある。"""

    __slots__ = ("_p",)

    def __init__(self, raw):
        self._p = raw

    @property
    def text(self):
        return self._p.text

    @text.setter
    def text(self, value):
        self._p.text = value

    def replace(self, old, new):
        return self._p.replace(old, new)

    @property
    def runs(self):
        return [Run(r) for r in self._p.runs]

    @property
    def style(self):
        """段落のスタイル。字としても `.name` でも読めます"""
        n = self._p.style
        return None if n is None else StyleName(n, self._p, sid=self._p.style_id)

    @style.setter
    def style(self, value):
        # python-docx のスタイルの物("Heading 1" を .name に持つ)も、
        # ただの文字("heading1")も受ける
        self._p.style = str(getattr(value, "name", value))

    @property
    def align(self):
        return self._p.align

    @align.setter
    def align(self, value):
        self._p.align = value

    @property
    def alignment(self):
        # python-docx の名前。中身は align と同じ字
        return self._p.align

    @alignment.setter
    def alignment(self, value):
        self._p.align = _align_word(value) or "left"

    @property
    def paragraph_format(self):
        return ParagraphFormat(self._p)

    @property
    def comments(self):
        return [Comment(a, t, self) for a, t in self._p.comments]

    @property
    def contains_page_break(self):
        """この段落で改ページするか(python-docx と同じ)。

        **模型は改ページを段落の性質で持ちます**(`page_break_before`)。
        本家は「改ページの run が中にあるか」で見ますが、紙の上の意味は
        同じです。
        """
        return bool(self._p.page_break_before)

    @property
    def rendered_page_breaks(self):
        """組んだ結果の改ページ。**組んでみないと分かりません**ので、
        ここは指定した改ページだけを返します(本家も指定の分は返します)"""
        return [self] if self._p.page_break_before else []

    def add_comment(self, text, author=""):
        """この段落にコメントを付ける(段落単位 — 文中の範囲は持たない)。"""
        self._p.add_comment(text, author)

    @property
    def in_table(self):
        return self._p.in_table

    # ── python-docx の口(互換層)───────────────────────────────

    def clear(self):
        """字を消す。段落の性質(見出し・寄せ)と先頭 run の書式は残る
        (python-docx の定義と同じ)。返りは自分。"""
        self._p.text = ""
        return self

    def add_run(self, text="", style=None):
        """段落の末尾に run を継ぎ足す(python-docx と同じ口)。
        書式は末尾の run のものを継ぐ。style は styles にある文字スタイル
        (無い名前は断る — add_style で作ってから)。"""
        r = Run(self._p.add_run(text))
        if style is not None:
            r.style = style
        return r

    @property
    def hyperlinks(self):
        """この段落のリンク(本家と同じ口)。.text と .address を持つ。"""
        return [Hyperlink(t, u) for t, u in self._p.hyperlinks]

    def add_hyperlink(self, text, address):
        """段落の末尾にリンクを足す。書式は末尾の run を継ぐ。"""
        return Run(self._p.add_hyperlink(text, address))

    def insert_paragraph_before(self, text=None, style=None):
        """この段落の前に段落を差す(python-docx と同じ口)。
        手元の段落の物は位置で指しているので、差した後は引き直すこと。"""
        p = Paragraph(self._p.insert_paragraph_before(text or ""))
        if style is not None:
            p.style = style
        return p

    def iter_inner_content(self):
        """段落の中身を順に。いまは run だけ(リンクの読みはエンジンの
        「足す」待ち — 台帳の hyperlinks)。"""
        for r in self.runs:
            yield r

    def __repr__(self):
        return repr(self._p)


class Cell:
    """表のセル。中には段落が並んでいる。"""

    __slots__ = ("_c",)

    def __init__(self, raw):
        self._c = raw

    def merge(self, other):
        """このセルから相手のセルまでを1つに結合する(python-docx と同じ)"""
        return Cell(self._c.merge(other._c))

    @property
    def vertical_alignment(self):
        """セルの中の縦位置。"top" / "center" / "bottom"(docx の既定は上)"""
        return self._c.vertical_alignment

    @vertical_alignment.setter
    def vertical_alignment(self, v):
        self._c.vertical_alignment = _valign_word(v)

    @property
    def width(self):
        """セルの幅。docx は列で持つので、この列の幅です"""
        w = self._c.width
        return Length.from_mm(w) if w else None

    @width.setter
    def width(self, v):
        self._c.width = 0.0 if v is None else _to_mm(v)

    @property
    def text(self):
        return self._c.text

    @text.setter
    def text(self, value):
        self._c.text = value

    @property
    def paragraphs(self):
        return [Paragraph(p) for p in self._c.paragraphs]

    def __repr__(self):
        return repr(self._c)


def _valign_word(v):
    """縦位置の言い方をそろえる。

    python-docx は `WD_ALIGN_VERTICAL.CENTER` のような列挙で渡します。
    名前でも字でも受けて、エンジンの言葉に直します。
    """
    if v is None:
        return "top"
    t = str(getattr(v, "name", v)).strip().lower()
    return {"middle": "center", "both": "center"}.get(t, t)


class _Settings:
    """`d.settings` の役。**中身は文書が直に持ちます**"""

    __slots__ = ("_d",)

    def __init__(self, doc):
        self._d = doc

    @property
    def odd_and_even_pages_header_footer(self):
        """奇数と偶数でヘッダーを分けているか。**文書の粒度では持ちません**
        (表の側は持ちます)ので False"""
        return False

    def __repr__(self):
        return "<Settings>"


class _HeadFoot(str):
    """ヘッダー / フッター。**字そのものとして振る舞います。**

    `print(d.header)` で中身が出て、`d.header.text = "…"` でも書けます。
    python-docx は段落の並びを返しますが、こちらは1行の字で足ります
    (中で改行すれば段落が分かれます)。
    """

    def __new__(cls, raw, which):
        self = super().__new__(cls, getattr(raw, which))
        self._raw = raw
        self._which = which
        return self

    @property
    def text(self):
        return str(self)

    @text.setter
    def text(self, v):
        setattr(self._raw, self._which, "" if v is None else str(v))

    def add_paragraph(self, text=""):
        """ヘッダー / フッターに段落を足します(python-docx と同じ口)。

        返りは**普通の段落**なので、揃えも書式も掛けられます。
        """
        return Paragraph(
            self._raw.add_hf_paragraph(text, footer=self._which == "footer"))

    @property
    def paragraphs(self):
        """段落の一覧(python-docx と同じ)。

        **段落が1つも無ければ空の段落を1つ作ります。** python-docx の
        新しい文書はヘッダーに空の段落を1つ持っていて、見本は
        `header.paragraphs[0]` から書き始めます(2026-09-01)。
        """
        ashi = self._which == "footer"
        ps = self._raw.hf_paragraphs(footer=ashi)
        if not ps:
            self._raw.add_hf_paragraph("", footer=ashi)
            ps = self._raw.hf_paragraphs(footer=ashi)
        return [Paragraph(p) for p in ps]


class Row:
    """表の1行。"""

    __slots__ = ("_row",)

    def __init__(self, raw):
        self._row = raw

    def __len__(self):
        return len(self._row)

    def __getitem__(self, i):
        return Cell(self._row[i])

    @property
    def height(self):
        """行の高さ。指定なしは None(python-docx と同じ)"""
        h = self._row.height
        return Length.from_mm(h) if h else None

    @height.setter
    def height(self, v):
        self._row.height = 0.0 if v is None else _to_mm(v)

    @property
    def cells(self):
        return [Cell(c) for c in self._row.cells]

    def __repr__(self):
        return repr(self._row)


class _Column:
    """表の1列。python-docx の columns[j].cells の形だけを持つ。"""

    __slots__ = ("_t", "_col")

    def __init__(self, table, col):
        self._t = table
        self._col = col

    @property
    def cells(self):
        return self._t.column_cells(self._col)

    @property
    def _raw(self):
        return self._t._t

    @property
    def width(self):
        """列の幅。**Length で返します**(`.mm` で mm になります)。
        指定していなければ `None`(等分)。"""
        w = self._raw.col_widths_mm
        return Length.from_mm(w[self._col]) if self._col < len(w) else None

    @width.setter
    def width(self, v):
        w = list(self._raw.col_widths_mm)
        # **足りない分は等分のまま**にはできません(mm の並びで持つので)。
        # 幅の分かっていない列は、A4 の本文の幅を等分した値で埋めます
        _, n = self._t.shape
        if len(w) < n:
            share = (210.0 - 40.0) / max(n, 1)
            w += [share] * (n - len(w))
        w[self._col] = _to_mm(v)
        self._raw.col_widths_mm = w

    def __repr__(self):
        return "<officework.doc 列 {}>".format(self._col)


class Table:
    """表。`t[行][列]` でセルに届く。"""

    __slots__ = ("_t",)

    def __init__(self, raw):
        self._t = raw

    @property
    def table(self):
        """自分自身(python-docx と同じ。セルから表へ辿るときの口)"""
        return self

    @property
    def table_direction(self):
        """表の向き。**右から左に並べる表は模型に持ちません**ので None。
        文書ぜんたいの右横書きは `ws.rtl`(表計算)側にあります。"""
        return None

    @table_direction.setter
    def table_direction(self, v):
        if v is not None:
            raise NotImplementedError(
                "右から左に並べる表は模型に持ちません(黙って受けません)"
            )

    def __len__(self):
        return len(self._t)

    def __getitem__(self, i):
        return Row(self._t[i])

    @property
    def rows(self):
        return [Row(r) for r in self._t.rows]

    @property
    def shape(self):
        return self._t.shape

    def values(self):
        return self._t.values()

    # ── python-docx の口(互換層)───────────────────────────────

    def add_row(self):
        """行を1つ足す(末尾)。明細行の継ぎ足し。返りは新しい行。"""
        return Row(self._t.add_row())

    def add_column(self, width=None):
        """列を1つ足す(右端)。width は python-docx と同じ EMU
        (docx.shared.Mm(25) 等がそのまま通る)。省略なら等分のまま。
        返りは新しい列。"""
        width_mm = None if width is None else width / 36000
        self._t.add_column(width_mm)
        return self.columns[-1]

    @property
    def style(self):
        """表のスタイルの名前(styleId)。定義は持たない — 名前を運ぶだけ。"""
        return self._t.style

    @style.setter
    def style(self, v):
        if v is None:
            self._t.style = None
            return
        # python-docx のスタイルの物(.style_id)も、名前の文字も受ける。
        # 名前("Table Grid")は Word の流儀で styleId("TableGrid")に寄せる
        sid = getattr(v, "style_id", None)
        if sid is None:
            sid = str(getattr(v, "name", v)).replace(" ", "")
        self._t.style = sid

    @property
    def alignment(self):
        return self._t.alignment

    @alignment.setter
    def alignment(self, v):
        self._t.alignment = _align_word(v)

    @property
    def autofit(self):
        return self._t.autofit

    @autofit.setter
    def autofit(self, v):
        self._t.autofit = bool(v)

    def cell(self, row_idx, col_idx):
        return Cell(self._t[row_idx][col_idx])

    def row_cells(self, row_idx):
        return Row(self._t[row_idx]).cells

    def column_cells(self, col_idx):
        # 結合のある帳票では行によって列数が違う。無い行は飛ばす
        # (python-docx は長方形しか持てないので、この場合の定義が向こうに無い)
        out = []
        for row in self._t.rows:
            if col_idx < len(row):
                out.append(Cell(row[col_idx]))
        return out

    @property
    def columns(self):
        _, cols = self._t.shape
        return [_Column(self, j) for j in range(cols)]

    def __repr__(self):
        return repr(self._t)


class Doc(NoStrayAttributes):
    """docx の文書。エンジンの Doc を包み、python-docx の口を足す。"""

    # 自分で持つ属性。ここに無い名前への代入は断ります(打ち間違い避け)
    _own = ("_d", "_path")

    def __init__(self, path=None, lang=None):
        """`Doc()` は空の文書、`Doc("報告.docx")` は開きます。

        python-docx の `Document(径路)` と同じ形です。前は `Doc.open` しか
        無く、本家の台本が1行目で止まりました(2026-08-28)。

        ``lang`` は組むときの言語です(``"ja"``, ``"en"`` など)。渡さない
        ときは、設定ファイルと OS の言語から決めます(2026-08-30)。
        """
        if path is not None:
            self._d = _doc.Doc.open(str(path), lang)
        else:
            self._d = _doc.Doc(lang)
        # 開いた元の名前。`to_pdf()` を引数なしで呼ぶときに使います
        self._path = str(path) if path is not None else None

    @staticmethod
    def open(path, lang=None):
        # **pathlib.Path も受ける**(python-docx と同じ。2026-08-15)。
        # 芯は文字しか取らないので、ここで径路の形に直してから渡す。
        # sheet.Book と揃えること — 片方だけ受けるのがいちばん困る
        d = Doc.__new__(Doc)
        d._path = _os.fspath(path)
        d._d = _doc.Doc.open(d._path, lang)
        return d

    def add_shape(self, kind, x, y, width, height, **kw):
        """**ページに貼り付く図形を置く。**

        置き場と大きさは紙の左上からの mm です。形の名前は xlsx と同じで、
        ``rect`` / ``roundRect`` / ``ellipse`` / ``rightArrow`` /
        ``diamond`` / ``line`` が使えます。

        ``fill`` と ``line`` は色("DDE7F0")、``text`` は図形の中の文字、
        ``rotation`` は回す角度、``opacity`` は不透明度、``shadow`` は影、
        ``page`` は何ページ目か(0始まり)です。

        python-docx には無い口です(本家は画像しか置けません)。
        """
        return self._d.add_shape(kind, x, y, width, height, **kw)

    @property
    def shapes(self):
        """置いた図形の数"""
        return self._d.shapes

    def save(self, path, dpi=None):
        """保存する。拡張子で行き先が決まります。

        ``.docx`` は文書、``.pdf`` は紙、``.png`` は絵です。
        ``dpi`` は絵の細かさで、既定は 150 です(``.png`` のときだけ効きます)。
        頁が複数あるときは、2枚目から名前に ``-2``・``-3`` が付きます。
        """
        # Path も受ける(上の open と同じ理由)
        self._d.save(_os.fspath(path), dpi)

    def fill(self, name, value):
        """名前の付いた記入欄すべてに値を入れます。返り値は入れた欄の数です。

        名前は writer の「フォーム」タブで記入欄に付けた物です。同じ名前の
        欄が2つあれば2つとも入ります。その名前の欄が無ければ KeyError で
        止まります(黙って空振りにしません)。本文の ``{{名前}}`` を埋めるのは
        `render` です。
        """
        n = self._d.fill(str(name), str(value))
        if n == 0:
            raise KeyError("記入欄「{}」がありません(d.fields() で名前を確かめられます)".format(name))
        return n

    def render(self, values, rows=None):
        """本文の ``{{名前}}`` に値を差し込みます。

        ``values`` は ``{"名前": 値}`` です。``rows`` は ``{"群": [行, …]}`` で、
        ``{{群.項目}}`` を含む表の行が、行の数だけ増えます。値は文字に直して
        入れます(数を渡してもかまいません)。データに無い名前は
        ``{{名前}}`` のまま残り、返り値の文にその名前が出ます。
        """
        moji = lambda v: "" if v is None else str(v)  # noqa: E731
        vals = {str(k): moji(v) for k, v in dict(values or {}).items()}
        gun = None
        if rows:
            gun = {str(g): [{str(k): moji(v) for k, v in dict(row).items()} for row in r]
                   for g, r in dict(rows).items()}
        return self._d.render(vals, gun)

    def page_count(self):
        """ページ数を数えます。PDF と同じ組み方で紙面を組みます(PDF は書きません)。"""
        return self._d.page_count()

    def to_pdf(self, path=None):
        """文書を PDF にします。返り値は保存先です。

        ``save("x.pdf")`` と同じ道を通ります。名前が違うだけではなく、
        PDF だけの指定を足すならこちらに足します(2026-08-30 発注者)。

        ``path`` を省くと、開いたファイルの名前の拡張子を ``.pdf`` に
        替えた所へ書きます。新しい文書なら ``文書1.pdf`` です。
        """
        if path is None:
            moto = getattr(self, "_path", None) or "文書1"
            path = _os.path.splitext(moto)[0] + ".pdf"
        path = _os.path.abspath(_os.fspath(path))
        self._d.save(path, None)
        return path

    @property
    def unsupported(self):
        return self._d.unsupported

    @property
    def paragraphs(self):
        return [Paragraph(p) for p in self._d.paragraphs]

    @property
    def tables(self):
        return [Table(t) for t in self._d.tables]

    def __getitem__(self, i):
        return Paragraph(self._d[i])

    def __len__(self):
        return len(self._d)

    @property
    def text(self):
        return self._d.text

    def add_comment(self, text, author="", paragraph=None):
        """**コメントを付ける**(python-docx と同じ口)。

        付ける先は段落です(模型の粒度)。`paragraph` を渡さなければ
        いちばん後ろの段落に付きます。
        """
        p = paragraph if paragraph is not None else self[len(self) - 1]
        p.add_comment(text, author)
        return p

    @property
    def settings(self):
        """文書の設定(python-docx の `settings` の役)。

        **こちらは設定を文書が直に持ちます** — 書体は `d.font`、
        字の大きさは `d.size_pt`、紙は `d.sections[0]` です。ここは
        本家の台本が読んでも落ちないよう、その入り口を返します。
        """
        return _Settings(self)

    @property
    def font(self):
        """文書ぜんたいの書体。PDF にもここが効きます"""
        return self._d.font

    @font.setter
    def font(self, v):
        self._d.font = v

    @property
    def size_pt(self):
        """文書ぜんたいの字の大きさ(pt)"""
        return self._d.size_pt

    @size_pt.setter
    def size_pt(self, v):
        self._d.size_pt = v

    @property
    def header(self):
        """ヘッダー。`d.header.text = "…"` でも `d.header = "…"` でも書けます"""
        return _HeadFoot(self._d, "header")

    @header.setter
    def header(self, v):
        self._d.header = "" if v is None else str(v)

    @property
    def footer(self):
        """フッター。ページ番号は `#`、総ページ数は `##` で書きます"""
        return _HeadFoot(self._d, "footer")

    @footer.setter
    def footer(self, v):
        self._d.footer = "" if v is None else str(v)

    def find(self, needle):
        return [Paragraph(p) for p in self._d.find(needle)]

    def replace(self, old, new):
        return self._d.replace(old, new)

    def add_paragraph(self, text="", style=None):
        """段落を足す(python-docx と同じ口)。style は名前でも様式の物でも。
        無い様式は**黙って作らない** — add_style で作ってから(家の作法)。"""
        p = Paragraph(self._d.add_paragraph(text))
        if style is not None:
            p.style = style
        return p

    def add_heading(self, text="", level=1):
        """見出しを足す(python-docx と同じ呼び方)。level は 0〜3 で、
        0 は文書の表題です。見出しは3段までなので、4 以上は断ります。"""
        return Paragraph(self._d.add_heading(text, level))

    def add_page_break(self):
        """改ページを足す(python-docx と同じ口)。本家は「改ページの run」を
        足すが、うちは**段落の性質(page_break_before)**で持つ — 紙の上の
        意味は同じで、本家の paragraph_format.page_break_before でも読める。"""
        return Paragraph(self._d.add_page_break())

    def add_picture(self, image, width=None, height=None):
        """画像を足す(python-docx と同じ口)。径路でも bytes でも。
        大きさは mm の数でも、本家の Length(Mm(60) 等)でもよい。
        返りは画像を持つ段落(本家は InlineShape — そこだけ流儀が違う)。"""
        def _mm(v):
            if v is None:
                return None
            return float(v.mm) if hasattr(v, "mm") else float(v)

        return Paragraph(self._d.add_picture(image, _mm(width), _mm(height)))

    def iter_inner_content(self):
        """段落と表を**文書の順**で返す(python-docx と同じ口)。"""
        for b in self._d.iter_inner_content():
            if isinstance(b, _doc.Table):
                yield Table(b)
            else:
                yield Paragraph(b)

    @property
    def core_properties(self):
        """文書の情報(author / title / keywords / subject / comments)。
        読み書きとも本家と同じ呼び名(author = docx の dc:creator)。"""
        return self._d.core_properties

    @property
    def styles(self):
        """スタイルの一覧(本家と同じ口 — 名前で引け、add_style で足せる)。
        定義の本体は styles.xml が持ち、保存で原本のまま持ち越される。"""
        return _Styles(self._d)

    @property
    def sections(self):
        """節の一覧(本家と同じ口)。途中の節+文書末の節。"""
        return [Section(s, self) for s in self._d.sections]

    def add_section(self, start_type=None):
        """節を足す(python-docx と同じ切り方)。**切るのは末尾** —
        いままで書いた分が前の節になり、これから足す物が新しい節に入る。
        新しい節は同じ紙と余白を継ぐので、変えるなら返ってきた節に書く。

        start_type は本家の WD_SECTION でも "new_page" / "continuous" でも。
        新しい段・偶数頁・奇数頁は模型に無いので正直に断る。"""
        kind = "new_page"
        if start_type is not None:
            name = getattr(start_type, "name", None)
            kind = str(name if name is not None else start_type).lower()
        return Section(self._d.add_section(kind), self)

    @property
    def inline_shapes(self):
        """文書の画像の一覧(本家と同じ口。width / height は Length)。
        本文の段落の分 — 表のセルの中の画像は数えない(模型の粒度)。"""
        out = []
        for p in self._d.paragraphs:
            for w, h in p.images:
                out.append(InlineShape(w, h))
        return out

    @property
    def comments(self):
        """文書のコメントの一覧(本家と同じ口)。うちは**段落単位**の粒度 —
        Comment.paragraph でどの段落かが分かる。付けるのは
        Paragraph.add_comment から。"""
        out = []
        for p in self.paragraphs:
            out.extend(p.comments)
        return out

    def add_table(self, rows, cols, style=None):
        """表を新しく組む(明細の帳票づくり)。各セルは空の段落を1つ持つ。

        `style` は**名前を運ぶだけ**です。定義(styles.xml)はこちらでは
        持たず、原本(雛形)が持っている前提です。名前を運べば Word で
        開いたときにその見た目になります。**組む所は名前を見ません**ので、
        officework の画面と PDF では罫線も帯も付きません。
        """
        t = Table(self._d.add_table(rows, cols))
        if style is not None:
            t._t.style = str(getattr(style, "name", style))
        return t

    def __getattr__(self, name):
        # エンジンに後から生えた口は、包み直しを待たずにそのまま通す
        if name.startswith("_"):  # 自分の畑(_d 等)で再帰しない
            raise AttributeError(name)
        return getattr(self._d, name)

    def __repr__(self):
        return repr(self._d)


__all__ = ["Doc", "Paragraph", "Run", "Table", "Row", "Cell"]
