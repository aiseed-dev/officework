# 互換層の適合検査 — **書きは定義どおり動作するか**(2026-08-12 発注者確定の合否線)。
#
# 本家(openpyxl / python-docx)が .venv に居れば、同じ手順を両方で動かして
# **結果そのものを突き合わせる**。居なければその節は飛ばす(無いのに失敗と言わない)。
# xlwings は Excel が無いと動かないので、参照の算術は文書の定義値と照合する。
#
# 手で回すなら:
#   .venv/bin/python pysheet/test_gokan.py
import os
import sys
import tempfile

from officework import sheet as office_sheet


def check(cond, msg):
    if not cond:
        print(f"NG: {msg}", file=sys.stderr)
        sys.exit(1)


# =============================================================== openpyxl の口
try:
    import openpyxl
except ImportError:
    openpyxl = None
    print("openpyxl が無いので突き合わせは飛ばした", file=sys.stderr)

b = office_sheet.Book()
s = b[0]

# --- 空のシートの端(openpyxl の定義: 空でも 1 と "A1:A1")---------------------
check(s.max_row == 1 and s.max_column == 1, f"空の max が 1 でない: {s.max_row},{s.max_column}")
check(s.min_row == 1 and s.min_column == 1, "空の min が 1 でない")
check(s.dimensions == "A1:A1", f"空の dimensions: {s.dimensions}")

# --- append: 使われている範囲の次の行へ。dict は列を選んで書く ------------------
s.append([1, 2])
s.append({"B": 9})          # 列の字で
s.append({3: "う"})         # 列の番号(1起点)でも
check(s["A1"] == 1 and s["B1"] == 2, f"append の1行目: {s['A1']},{s['B1']}")
check(s["B2"] == 9, f"append の dict(字): {s['B2']}")
check(s["C3"] == "う", f"append の dict(番号): {s['C3']}")
check(s.dimensions == "A1:C3", f"append 後の dimensions: {s.dimensions}")

# --- cell(row, column, value=): 書いて、札(座標)が openpyxl と同じ形 ----------
c = s.cell(row=2, column=3, value=7)
check(s["C2"] == 7, "cell(value=) が書けていない")
check(c.coordinate == "C2" and c.column_letter == "C" and c.col_idx == 3,
      f"Cell の札: {c.coordinate},{c.column_letter},{c.col_idx}")
check(c.data_type == "n", f"数の data_type: {c.data_type}")
check(c.offset(row=1, column=-2).coordinate == "A3", "offset の算術")
check(c.parent is s, "Cell.parent がシートに戻らない")
s["D1"] = "=A1+B1"
check(s.cell(row=1, column=4).data_type == "f", "式の data_type が 'f' でない")
check(s.cell(row=3, column=3).data_type == "s", "字の data_type が 's' でない")

# --- iter_rows / iter_cols / rows / columns ------------------------------------
got = [[cc.coordinate for cc in row] for row in s.iter_rows(min_row=1, max_row=2, max_col=2)]
check(got == [["A1", "B1"], ["A2", "B2"]], f"iter_rows の座標: {got}")
got = list(s.iter_rows(min_row=1, max_row=1, max_col=2, values_only=True))
check(got == [(1, 2)], f"iter_rows(values_only): {got}")
got = [[cc.coordinate for cc in col] for col in s.iter_cols(min_col=2, max_col=3, max_row=2)]
check(got == [["B1", "B2"], ["C1", "C2"]], f"iter_cols の座標: {got}")
check(len(list(s.rows)) == s.max_row, "rows が max_row と食い違う")
check(len(list(s.columns)) == s.max_column, "columns が max_column と食い違う")

# --- Workbook の口 --------------------------------------------------------------
check(b.sheetnames == b.sheet_names, "sheetnames が sheet_names と食い違う")
ws2 = b.create_sheet()
check(ws2.title in b.sheetnames, "create_sheet の既定の名前が一覧に無い")
check(b.index(ws2) == len(b) - 1, "index(ws) が末尾を指さない")
check(ws2.parent is b, "parent がブックに戻らない")
check([w.title for w in b.worksheets] == b.sheetnames, "worksheets の並び")
check(b.active.title == b.sheetnames[0], "active が先頭でない")
check(ws2.title in b and "居ない名前" not in b, "in(__contains__)が効かない")
b.close()  # 何もしないが、呼べること

# --- insert_rows / delete_rows / insert_cols / delete_cols(amount つき)--------
# openpyxl と同じ手順を並べて動かし、**盤面の値が同じになるか**を見る
if openpyxl is not None:
    def grid(rows, cols, at):
        return [[at(i, j) for j in range(1, cols + 1)] for i in range(1, rows + 1)]

    wb_o = openpyxl.Workbook()
    ws_o = wb_o.active
    b_j = office_sheet.Book()
    ws_j = b_j[0]
    for i in range(1, 4):
        for j in range(1, 4):
            ws_o.cell(row=i, column=j, value=i * 10 + j)
            ws_j.cell(row=i, column=j, value=i * 10 + j)
    for ws in (ws_o, ws_j):
        ws.insert_rows(2, amount=2)
        ws.delete_rows(1)
        ws.insert_cols(2, amount=1)
        ws.delete_cols(4, amount=2)
    g_o = grid(4, 2, lambda i, j: ws_o.cell(row=i, column=j).value)
    g_j = grid(4, 2, lambda i, j: ws_j.cell(row=i, column=j).value)
    check(g_o == g_j, f"行・列の出し入れが openpyxl と食い違う:\n  本家 {g_o}\n  うち {g_j}")

    # --- 定数(19個)は openpyxl の実物と同じ値 --------------------------------
    from openpyxl.worksheet.worksheet import Worksheet as _W
    for n in dir(_W):
        if n.split("_")[0] in ("BREAK", "ORIENTATION", "PAPERSIZE", "SHEETSTATE"):
            check(getattr(office_sheet.Sheet, n) == getattr(_W, n),
                  f"定数 {n} が openpyxl と違う")

    # --- 書いた物を本家が読めるか(定義どおりの何よりの証拠)--------------------
    with tempfile.TemporaryDirectory() as t:
        out = os.path.join(t, "gokan.xlsx")
        b2 = office_sheet.Book()
        s2 = b2[0]
        s2.append(["品名", "数", "単価", "金額"])
        s2.append(["ザボガードF", 4, 125000, "=B2*C2"])
        b2.save(out)
        r = openpyxl.load_workbook(out)
        rs = r.active
        check(rs["A1"].value == "品名" and rs["B2"].value == 4,
              "うちが書いた値を openpyxl が読めない")
        check(rs["D2"].value == "=B2*C2", "うちが書いた式を openpyxl が読めない")
        rv = openpyxl.load_workbook(out, data_only=True).active
        check(rv["D2"].value == 500000,
              f"うちが書き込んだ計算済みの値(openpyxl 自身は作れない物): {rv['D2'].value}")

        # **本家が作った字入りのブックを、うちが保存し直しても本家が読める。**
        # openpyxl の原本は共有文字列の関係を持たないので、持ち越しだけだと
        # 索引が外れて開けなくなっていた(2026-08-13 に踏んで直した)
        out_ss = os.path.join(t, "ss_roundtrip.xlsx")
        wb_ss = openpyxl.Workbook()
        ws_ss = wb_ss.active
        ws_ss.append(["品名", "金額"])
        ws_ss.append(["ザボガードF", 125000])
        wb_ss.save(out_ss)
        b_ss = office_sheet.Book.open(out_ss)
        out_ss2 = os.path.join(t, "ss_roundtrip_out.xlsx")
        b_ss.save(out_ss2)
        r_ss = openpyxl.load_workbook(out_ss2).active
        check(r_ss["A1"].value == "品名" and r_ss["A2"].value == "ザボガードF",
              f"字が往復しない(共有文字列の関係): {r_ss['A1'].value!r}")

        # 逆向き: 本家が書いた物をうちが読めるか
        out2 = os.path.join(t, "opx.xlsx")
        wb3 = openpyxl.Workbook()
        ws3 = wb3.active
        ws3.append([1, "あ", True])
        wb3.save(out2)
        b3 = office_sheet.Book.open(out2)
        s3 = b3[0]
        check(s3["A1"] == 1 and s3["B1"] == "あ" and s3["C1"] is True,
              "openpyxl が書いた物をうちが読めない")

# (第1歩では title の代入と create_sheet(index=) は「正直に断る」だったが、
#  第2歩でエンジンに書き口が入った — 下の第2歩の節で本式に検査する)

# --- 書式の書き: うちが書いた書式を本家が読める(定義どおりの何よりの証拠)------
if openpyxl is not None:
    with tempfile.TemporaryDirectory() as t:
        out = os.path.join(t, "fmt.xlsx")
        bf = office_sheet.Book()
        sf = bf[0]
        c = sf.cell(1, 1, value="題")
        c.font = office_sheet.Font(bold=True, size=14, color="FF0000")
        c.border = office_sheet.Border(
            top=office_sheet.Side("thin"),
            bottom=office_sheet.Side("double", office_sheet.Color("0070C0")),
        )
        c.fill = office_sheet.PatternFill("solid", fgColor="FFFF00")
        c.alignment = office_sheet.Alignment(
            horizontal="center", vertical="center", wrap_text=True
        )
        sf["B1"] = 45000
        sf.cell(1, 2).number_format = "yyyy/mm/dd"
        check(sf.cell(1, 2).is_date, "日付の表示形式なのに is_date が False")
        check(not sf.cell(1, 1).is_date, "字のセルまで is_date が True")
        bf.save(out)

        rc = openpyxl.load_workbook(out).active["A1"]
        check(rc.font.bold and rc.font.size == 14, f"font が本家で読めない: {rc.font}")
        check((rc.font.color.rgb or "").endswith("FF0000"), f"文字色: {rc.font.color}")
        check(rc.border.top.style == "thin" and rc.border.bottom.style == "double",
              f"罫線: {rc.border.top.style},{rc.border.bottom.style}")
        check((rc.border.bottom.color.rgb or "").endswith("0070C0"),
              f"罫線の色: {rc.border.bottom.color}")
        check(rc.fill.patternType == "solid"
              and (rc.fill.fgColor.rgb or "").endswith("FFFF00"),
              f"塗り: {rc.fill.patternType},{rc.fill.fgColor}")
        check(rc.alignment.horizontal == "center"
              and rc.alignment.vertical == "center" and rc.alignment.wrap_text,
              f"揃え: {rc.alignment}")
        rb = openpyxl.load_workbook(out).active["B1"]
        check(rb.number_format == "yyyy/mm/dd", f"表示形式: {rb.number_format}")

        # 逆向き: 本家が書いた書式をうちが読める
        out2 = os.path.join(t, "fmt_opx.xlsx")
        from openpyxl.styles import (Alignment as OAlign, Border as OBorder,
                                     Font as OFont, PatternFill as OFill,
                                     Side as OSide)
        wb4 = openpyxl.Workbook()
        ws4 = wb4.active
        oc = ws4["A1"]
        oc.value = "題"
        oc.font = OFont(bold=True, size=12, color="00B050")
        oc.border = OBorder(left=OSide(style="medium"))
        oc.fill = OFill("solid", fgColor="D9D9D9")
        oc.alignment = OAlign(horizontal="right", wrap_text=True)
        oc.number_format = "#,##0"
        wb4.save(out2)

        b5 = office_sheet.Book.open(out2)
        d = b5[0].fmt("A1")
        check(d.get("bold") and d.get("size") == 12.0, f"本家の font がうちで読めない: {d}")
        check(d.get("color") == "00B050", f"本家の文字色: {d.get('color')}")
        check(d.get("border_left", (None,))[0] == "medium", f"本家の罫線: {d}")
        check(d.get("fill") == "D9D9D9", f"本家の塗り: {d.get('fill')}")
        check(d.get("horizontal") == "right" and d.get("wrap"), f"本家の揃え: {d}")
        check(d.get("number_format") == "#,##0", f"本家の表示形式: {d}")
        # openpyxl の実物の入れ物をそのまま代入しても効く(属性名で受ける)
        c5 = b5[0].cell(2, 1)
        c5.font = OFont(italic=True)
        check(b5[0].fmt("A2").get("italic"), "openpyxl の Font の代入が効かない")

# --- セルの comment / hyperlink / protection(模型に既にある物の口)------------
if openpyxl is not None:
    with tempfile.TemporaryDirectory() as t:
        bc = office_sheet.Book()
        sc = bc[0]
        sc["A1"] = "確認"
        c1 = sc.cell(1, 1)
        c1.comment = "ここを見る"
        c1.hyperlink = "https://example.jp/"
        c1.protection = office_sheet.Protection(locked=False)
        check(c1.comment.text == "ここを見る", f"comment: {c1.comment}")
        check(c1.hyperlink.target == "https://example.jp/", f"hyperlink: {c1.hyperlink}")
        check(c1.protection.locked is False, f"protection: {c1.protection}")
        out_c = os.path.join(t, "cell_extras.xlsx")
        bc.save(out_c)

        rc2 = openpyxl.load_workbook(out_c).active["A1"]
        check(rc2.comment is not None and "ここを見る" in rc2.comment.text,
              f"うちのコメントを本家が読めない: {rc2.comment}")
        check(rc2.hyperlink is not None
              and rc2.hyperlink.target == "https://example.jp/",
              f"うちのリンクを本家が読めない: {rc2.hyperlink}")
        check(rc2.protection.locked is False, "うちの保護を本家が読めない")

        # 逆向き: 本家が書いた物をうちが読める
        from openpyxl.comments import Comment as OComment
        from openpyxl.styles import Protection as OProtection
        wb6 = openpyxl.Workbook()
        ws6 = wb6.active
        ws6["B2"] = 1
        ws6["B2"].comment = OComment("要確認", "甲")
        ws6["B2"].hyperlink = "https://example.jp/b2"
        ws6["B2"].protection = OProtection(locked=False)
        out_c2 = os.path.join(t, "cell_extras_opx.xlsx")
        wb6.save(out_c2)
        b7 = office_sheet.Book.open(out_c2)
        c7 = b7[0].cell(2, 2)
        check(c7.comment is not None and "要確認" in c7.comment.text,
              f"本家のコメント: {c7.comment}")
        check(c7.hyperlink is not None and c7.hyperlink.target.startswith("https://"),
              f"本家のリンク: {c7.hyperlink}")
        check(c7.protection.locked is False, "本家の保護")
        # 本家の Comment の物をそのまま代入しても効く
        c7.comment = OComment("直した", "乙")
        check(c7.comment.text == "直した", "本家の Comment の代入")

# --- 表(テーブル): 作れて・構造化参照が計算されて・本家と往復する ---------------
if openpyxl is not None:
    with tempfile.TemporaryDirectory() as t:
        bt = office_sheet.Book()
        st = bt[0]
        st.append(["品名", "金額"])
        st.append(["ザボガードF", 125000])
        st.append(["F-02", 225000])
        st.add_table(office_sheet.Table(
            displayName="明細", ref="A1:B3",
            tableStyleInfo=office_sheet.TableStyleInfo(name="TableStyleMedium9")))
        check("明細" in st.tables and st.tables["明細"].ref == "A1:B3",
              f"表の一覧: {dict(st.tables)}")
        # **構造化参照が計算まで効く**(openpyxl は式を計算しない = 上位分)
        st["D1"] = "=SUM(明細[金額])"
        check(st["D1"] == 350000, f"構造化参照の計算: {st['D1']}")
        out_t = os.path.join(t, "table.xlsx")
        bt.save(out_t)

        rt = openpyxl.load_workbook(out_t).active
        check("明細" in rt.tables, f"うちの表を本家が読めない: {list(rt.tables)}")
        check(rt.tables["明細"].ref == "A1:B3", f"表の範囲: {rt.tables['明細'].ref}")
        check(rt["D1"].value == "=SUM(明細[金額])", "構造化参照の式が往復しない")

        # 逆向き: 本家が作った表をうちが読み、式が計算できる
        from openpyxl.worksheet.table import Table as OTable, TableStyleInfo as OTSI
        wb_t2 = openpyxl.Workbook()
        ws_t2 = wb_t2.active
        ws_t2.append(["名", "数"])
        ws_t2.append(["甲", 3])
        ws_t2.append(["乙", 4])
        otb = OTable(displayName="在庫", ref="A1:B3")
        otb.tableStyleInfo = OTSI(name="TableStyleLight1", showRowStripes=True)
        ws_t2.add_table(otb)
        out_t2 = os.path.join(t, "table_opx.xlsx")
        wb_t2.save(out_t2)
        b_t2 = office_sheet.Book.open(out_t2)
        s_t2 = b_t2[0]
        check("在庫" in s_t2.tables, f"本家の表: {dict(s_t2.tables)}")
        check(s_t2.tables["在庫"].tableStyleInfo.name == "TableStyleLight1",
              "表の様式の名前が読めない")
        s_t2["D1"] = "=SUM(在庫[数])"
        check(s_t2["D1"] == 7, f"本家の表への構造化参照: {s_t2['D1']}")
        # 本家の実物の Table をそのまま渡しても効く
        s_t2.add_table(OTable(displayName="控え", ref="A1:B2"))
        check("控え" in s_t2.tables, "本家の Table の代入")
        s_t2.remove_table("控え")
        check("控え" not in s_t2.tables, "表が外れない")
        try:
            s_t2.add_table(office_sheet.Table(displayName="悪い 名前", ref="A1:B2"))
            check(False, "空白入りの名前が黙って通った")
        except ValueError:
            pass

# --- 名前付き範囲: 定義して・式で使えて・本家と往復する -------------------------
if openpyxl is not None:
    with tempfile.TemporaryDirectory() as t:
        bn = office_sheet.Book()
        sn = bn[0]
        sn["A1"] = 100
        sn["A2"] = 4
        bn.create_named_range("単価", bn.worksheets[0], "$A$1")
        bn.defined_names["数量"] = office_sheet.DefinedName(
            "数量", attr_text="{}!$A$2".format(sn.title))
        sn["B1"] = "=単価*数量"
        check(sn["B1"] == 400, f"名前が式で効かない: {sn['B1']}")
        check("単価" in bn.defined_names and len(bn.defined_names) == 2,
              f"defined_names: {dict(bn.defined_names)}")
        out_n = os.path.join(t, "names.xlsx")
        bn.save(out_n)

        rn = openpyxl.load_workbook(out_n)
        got = {k: v.attr_text.replace("$", "") for k, v in rn.defined_names.items()}
        check(got.get("単価", "").endswith("!A1"),
              f"うちの名前を本家が読めない: {got}")

        # 逆向き: 本家が定義した名前をうちが読み、式が計算される
        from openpyxl.workbook.defined_name import DefinedName as ODefinedName
        wb8 = openpyxl.Workbook()
        ws8 = wb8.active
        ws8["A1"] = 250
        wb8.defined_names["tanka"] = ODefinedName(
            "tanka", attr_text="{}!$A$1".format(ws8.title))
        out_n2 = os.path.join(t, "names_opx.xlsx")
        wb8.save(out_n2)
        b9 = office_sheet.Book.open(out_n2)
        s9 = b9[0]
        check("tanka" in b9.defined_names, f"本家の名前: {dict(b9.defined_names)}")
        s9["B1"] = "=tanka*2"
        check(s9["B1"] == 500, f"本家の名前が式で効かない: {s9['B1']}")
        del b9.defined_names["tanka"]
        check("tanka" not in b9.defined_names, "名前が消えない")

        # --- 印刷範囲と入力規則 ------------------------------------------------
        bp = office_sheet.Book()
        sp = bp[0]
        sp["A1"] = 1
        sp.print_area = "A1:C10"
        check(sp.print_area == "'{}'!$A$1:$C$10".format(sp.title),
              f"print_area の形: {sp.print_area}")
        dv = office_sheet.DataValidation(type="list", formula1='"甲,乙,丙"')
        dv.add("B1:B5")
        sp.add_data_validation(dv)
        check(sp.validations == [("B1:B5", "list", '"甲,乙,丙"', "", "")],
              f"validations: {sp.validations}")
        out_p = os.path.join(t, "print_dv.xlsx")
        bp.save(out_p)

        rp = openpyxl.load_workbook(out_p).active
        check(rp.print_area.replace("$", "").endswith("!A1:C10"),
              f"うちの印刷範囲を本家が読めない: {rp.print_area}")
        dvs = rp.data_validations.dataValidation
        check(len(dvs) == 1 and dvs[0].type == "list"
              and dvs[0].formula1 == '"甲,乙,丙"'
              and str(dvs[0].sqref) == "B1:B5",
              f"うちの入力規則を本家が読めない: {dvs}")

        # 印刷のタイトル行(頁ごとに繰り返す見出し — 複数頁の明細の定番)
        sp.print_title_rows = "1:2"
        check(sp.print_title_rows == "1:2", f"タイトル行: {sp.print_title_rows}")
        check(sp.print_titles == "'{}'!$1:$2".format(sp.title),
              f"print_titles の形: {sp.print_titles}")
        try:
            sp.print_title_cols = "A:B"
            check(False, "列の繰り返しが黙って通った")
        except NotImplementedError:
            pass
        out_pt = os.path.join(t, "titles.xlsx")
        bp.save(out_pt)
        rpt = openpyxl.load_workbook(out_pt).active
        check(rpt.print_title_rows == "$1:$2",
              f"うちのタイトル行を本家が読めない: {rpt.print_title_rows}")
        # 逆向き
        wb_t = openpyxl.Workbook()
        wb_t.active["A1"] = 1
        wb_t.active.print_title_rows = "1:3"
        out_pt2 = os.path.join(t, "titles_opx.xlsx")
        wb_t.save(out_pt2)
        check(office_sheet.Book.open(out_pt2)[0].print_title_rows == "1:3",
              "本家のタイトル行がうちで読めない")

        # 逆向き: 本家が書いた物をうちが読める(実物の DataValidation の代入も)
        from openpyxl.worksheet.datavalidation import DataValidation as ODV
        wb10 = openpyxl.Workbook()
        ws10 = wb10.active
        ws10.print_area = "B2:D4"
        odv = ODV(type="list", formula1="$D$2:$D$5")
        ws10.add_data_validation(odv)
        odv.add("A1:A3")
        out_p2 = os.path.join(t, "print_dv_opx.xlsx")
        wb10.save(out_p2)
        b11 = office_sheet.Book.open(out_p2)
        s11 = b11[0]
        check(s11.print_area.replace("$", "").endswith("!B2:D4"),
              f"本家の印刷範囲: {s11.print_area}")
        check(any(v[0] == "A1:A3" and v[1] == "list" for v in s11.validations),
              f"本家の入力規則: {s11.validations}")
        s11.add_data_validation(odv)  # 本家の実物をそのまま渡しても効く
        check(len(s11.validations) == 2, "本家の DataValidation の代入")

# --- 1904 起点: 読めて・計算と表示が正しく・datetime が往復する ----------------
if openpyxl is not None:
    import datetime

    with tempfile.TemporaryDirectory() as t:
        # 本家で 1904 起点のブックを作る
        wb12 = openpyxl.Workbook()
        wb12.epoch = openpyxl.utils.datetime.CALENDAR_MAC_1904
        ws12 = wb12.active
        ws12["A1"] = datetime.date(2026, 8, 13)
        ws12["A1"].number_format = "yyyy/m/d"
        out_e = os.path.join(t, "mac1904.xlsx")
        wb12.save(out_e)

        b13 = office_sheet.Book.open(out_e)
        s13 = b13[0]
        check(b13.epoch.year == 1904, f"起点の読み: {b13.epoch}")
        check(b13.excel_base_date.year == 1904, "excel_base_date(別名)")
        # 表示: 1899 起点で読むと 4 年ずれる — 正しく 2026 で出るか
        check(s13.display("A1").startswith("2026"),
              f"1904 起点の表示が4年ずれている: {s13.display('A1')}")
        # 関数: YEAR も起点どおり
        s13["B1"] = "=YEAR(A1)"
        check(s13["B1"] == 2026, f"1904 起点の YEAR: {s13['B1']}")
        # datetime の書き込みも起点どおりの通し番号になる
        s13["C1"] = datetime.date(2026, 8, 13)
        check(s13["C1"] == s13["A1"], f"datetime の受けが起点とずれる: "
              f"{s13['C1']} vs {s13['A1']}")
        # 往復して本家が同じ日付で読める
        out_e2 = os.path.join(t, "mac1904_rt.xlsx")
        b13.save(out_e2)
        r13 = openpyxl.load_workbook(out_e2)
        check(r13.epoch.year == 1904, "往復で起点が消えた")
        got = r13.active["A1"].value
        check(getattr(got, "year", None) == 2026 and got.month == 8 and got.day == 13,
              f"往復の日付を本家が読めない: {got!r}")

        # 普通のブック(1899)はそのまま
        b14 = office_sheet.Book()
        check(b14.epoch.year == 1899, f"既定の起点: {b14.epoch}")

# ================================================== xlwings の口(参照の算術)
# 橋は動いているアプリが要るので、ソケットに出ない算術だけを定義値と照合する
from officework import calc as xw

r = xw.Range("B2:D5")
check(r.address == "$B$2:$D$5", f"address: {r.address}")
check(r.get_address(False, False) == "B2:D5", "get_address(相対)")
check(r.row == 2 and r.column == 2, f"row,column: {r.row},{r.column}")
check(r.shape == (4, 3) and r.size == 12 and r.count == 12 and len(r) == 12,
      f"shape,size: {r.shape},{r.size}")
check(len(r.rows) == 4 and r.rows[0]._a1() == "B2:D2", "rows の刻み")
check(len(r.columns) == 3 and r.columns[2]._a1() == "D2:D5", "columns の刻み")
check(r.offset(1, 2)._a1() == "D3:F6", f"offset: {r.offset(1, 2)._a1()}")
check(r.resize(2, 2)._a1() == "B2:C3", f"resize: {r.resize(2, 2)._a1()}")
check(r.resize(row_size=1)._a1() == "B2:D2", "resize(片方だけ)")
check(r.last_cell._a1() == "D5", f"last_cell: {r.last_cell._a1()}")
one = xw.Range("A1")
check(one.address == "$A$1" and one.shape == (1, 1), "1マスの算術")
sh = xw.Sheet("見積")
check(sh.cells._a1() == "A1:XFD1048576", "Sheet.cells が全マスでない")
check(sh["B2"]._a1() == "B2", "Sheet の添字")

# =============================================== python-docx の口(表と段落)
try:
    import docx as pydocx
except ImportError:
    pydocx = None
    print("python-docx が無いので突き合わせは飛ばした", file=sys.stderr)

from officework import doc as office_doc

if pydocx is not None:
    with tempfile.TemporaryDirectory() as t:
        # 本家で作った文書を、うちで読んで**同じ添字が同じセルを指す**か
        src = os.path.join(t, "hon.docx")
        d_o = pydocx.Document()
        d_o.add_paragraph("最初の段落")
        tb = d_o.add_table(rows=3, cols=2)
        for i in range(3):
            for j in range(2):
                tb.cell(i, j).text = f"{i}-{j}"
        d_o.save(src)

        d_j = office_doc.Doc.open(src)
        t_j = d_j.tables[0]
        t_o = pydocx.Document(src).tables[0]
        for i in range(3):
            for j in range(2):
                check(t_j.cell(i, j).text == t_o.cell(i, j).text,
                      f"cell({i},{j}) が本家と食い違う")
        check([c.text for c in t_j.row_cells(1)] == [c.text for c in t_o.row_cells(1)],
              "row_cells が本家と食い違う")
        check([c.text for c in t_j.column_cells(1)] == [c.text for c in t_o.column_cells(1)],
              "column_cells が本家と食い違う")
        check(len(t_j.columns) == len(t_o.columns), "columns の数")
        check([c.text for c in t_j.columns[1].cells] == [c.text for c in t_o.columns[1].cells],
              "columns[j].cells が本家と食い違う")

        # clear: 字は消え、段落の性質は残り、自分が返る(本家と同じ定義)
        p_o = pydocx.Document(src).paragraphs[0]
        ret_o = p_o.clear()
        p_j = d_j[0]
        ret_j = p_j.clear()
        check(p_j.text == "" and p_o.text == "", "clear で字が消えない")
        check(ret_j is p_j and ret_o is p_o, "clear が自分を返さない")

        # iter_inner_content: run が順に出る
        d_j2 = office_doc.Doc.open(src)
        p = d_j2[0]
        check([r.text for r in p.iter_inner_content()] == [r.text for r in p.runs],
              "iter_inner_content が runs と食い違う")

        # 書きの往復: うちがセルへ書いた物を本家が読めるか
        t_j2 = d_j2.tables[0]
        t_j2.cell(1, 1).text = "書き換えた"
        out = os.path.join(t, "kaki.docx")
        d_j2.save(out)
        back = pydocx.Document(out)
        check(back.tables[0].cell(1, 1).text == "書き換えた",
              "うちが書いたセルを本家が読めない")

        # font の両対応: 字の比べも .name も通る
        rr = d_j2[0].runs if d_j2[0].runs else None
        if rr:
            f = rr[0].font
            check(isinstance(f, str) and (f.name is None or isinstance(f.name, str)),
                  "font の両対応(str と .name)が崩れている")

        # --- 段落の書式: add_heading・style の書き・paragraph_format ----------
        # うちが書いた物を本家が読める(定義どおりの何よりの証拠)
        d_h = office_doc.Doc()
        h = d_h.add_heading("第1章 概要", level=1)
        check(h.style == "heading1", f"add_heading の役目: {h.style}")
        p_f = d_h.add_paragraph("本文です")
        p_f.paragraph_format.alignment = "center"
        p_f.paragraph_format.line_spacing = 1.5
        p_f.paragraph_format.page_break_before = True
        p2 = d_h.add_paragraph("次の段落")
        p2.style = "Heading 2"          # python-docx の名前でも受ける
        check(p2.style == "heading2", f"style の書き(本家の名前): {p2.style}")
        out_h = os.path.join(t, "heading.docx")
        d_h.save(out_h)

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        back_h = pydocx.Document(out_h)
        check(back_h.paragraphs[0].style.name == "Heading 1",
              f"うちの見出しを本家が読めない: {back_h.paragraphs[0].style.name}")
        bf = back_h.paragraphs[1].paragraph_format
        check(bf.alignment == WD_ALIGN_PARAGRAPH.CENTER,
              f"うちの寄せを本家が読めない: {bf.alignment}")
        check(bf.line_spacing == 1.5, f"うちの行間を本家が読めない: {bf.line_spacing}")
        check(bf.page_break_before, "うちの改ページ前を本家が読めない")
        check(back_h.paragraphs[2].style.name == "Heading 2",
              "style の書きが保存で消えた")

        # 逆向き: 本家が書いた物をうちが読める
        d_o2 = pydocx.Document()
        d_o2.add_heading("題", level=2)
        po = d_o2.add_paragraph("中央寄せ")
        po.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        po.paragraph_format.line_spacing = 2.0
        out_o = os.path.join(t, "heading_opx.docx")
        d_o2.save(out_o)
        d_r = office_doc.Doc.open(out_o)
        check(d_r[0].style == "heading2", f"本家の見出しがうちで読めない: {d_r[0].style}")
        check(d_r[1].paragraph_format.alignment == "center",
              f"本家の寄せ: {d_r[1].paragraph_format.alignment}")
        check(d_r[1].paragraph_format.line_spacing == 2.0,
              f"本家の行間: {d_r[1].paragraph_format.line_spacing}")
        # 本家の enum をそのまま代入しても効く
        d_r[1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        check(d_r[1].align == "right", "本家の enum の代入が効かない")
        # 模型に無い物は黙って捨てない
        try:
            d_r[1].paragraph_format.space_before = 12
            check(False, "space_before が黙って通った")
        except NotImplementedError:
            pass
        try:
            d_h.add_heading("題", level=0)
            check(False, "level=0(Title)が黙って通った")
        except ValueError:
            pass

        # --- 表の書式: style(名前だけ運ぶ)・alignment・autofit ---------------
        from docx.enum.table import WD_TABLE_ALIGNMENT
        d_t = pydocx.Document()
        tb2 = d_t.add_table(rows=2, cols=2, style="Table Grid")
        tb2.alignment = WD_TABLE_ALIGNMENT.CENTER
        tb2.autofit = False
        out_t = os.path.join(t, "tbl.docx")
        d_t.save(out_t)

        d_tr = office_doc.Doc.open(out_t)
        t_r = d_tr.tables[0]
        check(t_r.style == "TableGrid", f"本家の表スタイルがうちで読めない: {t_r.style}")
        check(t_r.alignment == "center", f"本家の表の置き方: {t_r.alignment}")
        check(t_r.autofit is False, f"本家の autofit: {t_r.autofit}")

        # うちが書き替えた物を本家が読める(スタイル定義は原本の物が持ち越される)
        t_r.alignment = WD_TABLE_ALIGNMENT.RIGHT  # 本家の enum をそのまま
        t_r.autofit = True
        out_t2 = os.path.join(t, "tbl2.docx")
        d_tr.save(out_t2)
        back_t = pydocx.Document(out_t2).tables[0]
        check(back_t.style.name == "Table Grid",
              f"うちが運んだスタイル名を本家が読めない: {back_t.style.name}")
        check(back_t.alignment == WD_TABLE_ALIGNMENT.RIGHT,
              f"うちの表の置き方を本家が読めない: {back_t.alignment}")
        check(back_t.autofit, "うちの autofit を本家が読めない")

        # --- run の手: add_run・性質ごとの書き・add_text・clear ----------------
        d_run = office_doc.Doc()
        pr = d_run.add_paragraph("請求先: ")
        r2 = pr.add_run("株式会社甲")
        r2.bold = True
        r2.font.size = 12
        r2.add_text(" 御中")
        check(pr.text == "請求先: 株式会社甲 御中", f"add_run/add_text: {pr.text}")
        check(pr.runs[1].bold and pr.runs[1].size_pt == 12, "run の性質の書き")
        d_run.add_page_break()
        p_last = d_run.add_paragraph("次の頁")
        out_r = os.path.join(t, "runs.docx")
        d_run.save(out_r)

        back_r = pydocx.Document(out_r)
        rr2 = back_r.paragraphs[0].runs
        check(len(rr2) == 2 and rr2[1].bold and rr2[1].font.size.pt == 12,
              f"うちの run の書式を本家が読めない: {[(r.text, r.bold) for r in rr2]}")
        check(back_r.paragraphs[0].text == "請求先: 株式会社甲 御中",
              "run の字が本家で崩れる")
        # 改ページはうちの流儀(page_break_before)— 本家でもその形で読める
        check(back_r.paragraphs[1].paragraph_format.page_break_before,
              "add_page_break が本家で読めない")

        # 本家の run の作法(clear が自分を返す・add_text が書式を保つ)と同じか
        d_o3 = pydocx.Document()
        po3 = d_o3.add_paragraph("あ")
        ro3 = po3.add_run("い")
        ro3.bold = True
        ro3.add_text("う")
        check(po3.text == "あいう" and po3.runs[1].bold, "(本家の前提の確認)")
        r_mine = office_doc.Doc().add_paragraph("あ").add_run("い")
        r_mine.bold = True
        r_mine.add_text("う")
        check(r_mine.text == "いう" and r_mine.bold, "うちの add_text が書式を落とす")
        check(r_mine.clear() is r_mine and r_mine.text == "" and r_mine.bold,
              "clear が自分を返さない・書式まで消える")

        # --- 文書の順・途中に差す・文書の情報・画像 ----------------------------
        d_mix = office_doc.Doc()
        d_mix.add_paragraph("前")
        d_mix.add_table(1, 1)
        after = d_mix.add_paragraph("後")
        kinds = [type(x).__name__ for x in d_mix.iter_inner_content()]
        check(kinds == ["Paragraph", "Table", "Paragraph"],
              f"iter_inner_content の順: {kinds}")
        after.insert_paragraph_before("間")
        check([p.text for p in d_mix.paragraphs] == ["前", "間", "後"],
              f"insert_paragraph_before: {[p.text for p in d_mix.paragraphs]}")

        d_mix.core_properties.author = "日本不燃 太郎"
        d_mix.core_properties.title = "見積書"
        # 画像(2×2 の最小 PNG)を径路の代わりに bytes で
        png = (b"\x89PNG\r\n\x1a\n" +
               b"\x00\x00\x00\rIHDR\x00\x00\x00\x02\x00\x00\x00\x02"
               b"\x08\x02\x00\x00\x00\xfd\xd4\x9as" +
               b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05\xfe\x02\xfe"
               b"\xa75\x81\x84\x00\x00\x00\x00IEND\xaeB`\x82")
        d_mix.add_picture(png, width=30)  # mm。縦横比を保って 30×30
        out_m = os.path.join(t, "mix.docx")
        d_mix.save(out_m)

        back_m = pydocx.Document(out_m)
        check(back_m.core_properties.author == "日本不燃 太郎"
              and back_m.core_properties.title == "見積書",
              f"文書の情報を本家が読めない: {back_m.core_properties.author}")
        check(len(back_m.inline_shapes) == 1, "うちの画像を本家が読めない")
        check(round(back_m.inline_shapes[0].width.mm) == 30,
              f"画像の大きさ: {back_m.inline_shapes[0].width.mm}")
        # 逆向き: 本家の文書の情報をうちが読める
        d_o4 = pydocx.Document()
        d_o4.core_properties.author = "甲"
        d_o4.core_properties.comments = "控え"
        out_o4 = os.path.join(t, "props.docx")
        d_o4.save(out_o4)
        d_r4 = office_doc.Doc.open(out_o4)
        check(d_r4.core_properties.author == "甲"
              and d_r4.core_properties.comments == "控え",
              "本家の文書の情報がうちで読めない")

        # --- スタイル定義を運ぶ(2026-08-12 発注者確定「持たない主義では無理」)--
        # (1) 知らないスタイル名が保存で消えない — 「書式は据え置き」の穴を塞ぐ
        d_s = pydocx.Document()
        d_s.styles.add_style("社内様式", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        ps = d_s.add_paragraph("様式の段落")
        ps.style = d_s.styles["社内様式"]
        out_s = os.path.join(t, "style.docx")
        d_s.save(out_s)

        d_sr = office_doc.Doc.open(out_s)
        check("社内様式" in d_sr.styles, f"本家のスタイルの名乗り: {list(d_sr.styles)}")
        check(d_sr.styles["社内様式"].type == "paragraph", "スタイルの種類")
        check(d_sr[0].style == "社内様式", f"知らないスタイル名の読み: {d_sr[0].style}")
        out_s2 = os.path.join(t, "style_rt.docx")
        d_sr.save(out_s2)  # 開いて保存 — スタイルが消えないか
        back_s = pydocx.Document(out_s2)
        check(back_s.paragraphs[0].style.name == "社内様式",
              f"スタイル名が保存で消えた: {back_s.paragraphs[0].style.name}")

        # (2) うちで足して・張って、本家が読める
        d_a = office_doc.Doc()
        d_a.styles.add_style("引用風", "paragraph")
        d_a.styles.add_style("強調字", "character")
        check("引用風" in d_a.styles and len(d_a.styles) >= 6,  # 最小定義4+2
              f"add_style: {[s.name for s in d_a.styles]}")
        pa = d_a.add_paragraph("引用の段落")
        pa.style = "引用風"
        check(d_a[0].style == "引用風", f"段落スタイルの張り: {d_a[0].style}")
        ra = pa.add_run("大事", style="強調字")
        check(ra.style == "強調字", f"文字スタイルの張り: {ra.style}")
        try:
            pa.style = "無い様式"
            check(False, "無いスタイルが黙って通った")
        except ValueError:
            pass
        out_a = os.path.join(t, "style_add.docx")
        d_a.save(out_a)
        back_a = pydocx.Document(out_a)
        check(back_a.paragraphs[0].style.name == "引用風",
              f"うちの段落スタイルを本家が読めない: {back_a.paragraphs[0].style.name}")
        check(back_a.paragraphs[0].runs[-1].style.name == "強調字",
              f"うちの文字スタイルを本家が読めない: {back_a.paragraphs[0].runs[-1].style}")
        st_names = [s.name for s in back_a.styles]
        check("引用風" in st_names and "強調字" in st_names,
              f"うちの styles.xml 追記を本家が読めない: {st_names}")

        # --- 節(sections): 実物の複数節を読めて・余白の書きが原文に効く --------
        from docx.enum.section import WD_SECTION
        from docx.shared import Mm as OMm
        d_sec = pydocx.Document()
        d_sec.add_paragraph("1節目")
        s2_ = d_sec.add_section(WD_SECTION.NEW_PAGE)
        s2_.page_width = OMm(297)   # 2節目は A4 横
        s2_.page_height = OMm(210)
        d_sec.add_paragraph("2節目")
        out_sec = os.path.join(t, "sections.docx")
        d_sec.save(out_sec)

        d_sr2 = office_doc.Doc.open(out_sec)
        secs = d_sr2.sections
        check(len(secs) == 2, f"節の数: {len(secs)}")
        # 本家の新規文書の既定は Letter(216×279)— A4 ではない
        check(round(secs[0].page_width.mm) == 216 and secs[0].orientation == "portrait",
              f"1節目の紙: {secs[0]}")
        check(round(secs[1].page_width.mm) == 297 and secs[1].orientation == "landscape",
              f"2節目の紙: {secs[1]}")
        # 余白を書き替え → 保存 → 本家が読める(原文の sectPr への属性差し替え)
        secs[1].left_margin = OMm(30)
        check(round(secs[1].left_margin.mm) == 30, "余白の書きが読み戻せない")
        out_sec2 = os.path.join(t, "sections_rt.docx")
        d_sr2.save(out_sec2)
        back_sec = pydocx.Document(out_sec2)
        check(len(back_sec.sections) == 2, "往復で節が消えた")
        check(round(back_sec.sections[1].left_margin.mm) == 30,
              f"うちの余白を本家が読めない: {back_sec.sections[1].left_margin.mm}")
        check(round(back_sec.sections[1].page_width.mm) == 297,
              "差し替えで紙の大きさが崩れた")

        # --- inline_shapes(画像の読みの対)と コメント -------------------------
        d_m2 = office_doc.Doc.open(out_m)  # さっき画像を入れた文書
        shp = d_m2.inline_shapes
        check(len(shp) == 1 and round(shp[0].width.mm) == 30,
              f"inline_shapes: {shp}")
        p0 = d_m2[0]
        p0.add_comment("この行を確認", author="乙")
        check(d_m2.comments[0].text == "この行を確認"
              and d_m2.comments[0].author == "乙"
              and d_m2.comments[0].paragraph.text == p0.text,
              f"コメントの読み書き: {d_m2.comments}")
        out_c = os.path.join(t, "cmt.docx")
        d_m2.save(out_c)
        d_c = office_doc.Doc.open(out_c)
        check(d_c.comments and d_c.comments[0].text == "この行を確認",
              "コメントが保存で消えた")
        # 本家(1.2 以降)がコメント API を持つなら突き合わせる
        back_c = pydocx.Document(out_c)
        if hasattr(back_c, "comments"):
            check(any(c.text == "この行を確認" for c in back_c.comments),
                  f"うちのコメントを本家が読めない: {[c.text for c in back_c.comments]}")

# ==================== 第2歩(足すの背骨): 結合・固定枠・改名・複製・削除・並べ替え
b = office_sheet.Book()
s = b[0]

# 結合: 家の作法(アプリと同じ)— 左上以外の中身は消え、空の左上へは最初の中身が移る
s["B1"] = "題"
s["C2"] = 9
s.merge_cells("A1:C2")
check(s.merged_cell_ranges == ["A1:C2"], f"結合が台帳に載らない: {s.merged_cell_ranges}")
check(s["A1"] == "題", "空だった左上へ最初の中身が移っていない")
check(s["B1"] is None and s["C2"] is None, "呑まれた中身が消えていない")

# 解除: openpyxl と同じ定義 — その範囲そのものが結合でなければ ValueError
try:
    s.unmerge_cells("A1:B1")
    check(False, "結合でない範囲の解除が黙って通った")
except ValueError:
    pass
s.unmerge_cells("A1:C2")
check(s.merged_cell_ranges == [], "解除できていない")

# openpyxl の数字指定でも
s.merge_cells(start_row=1, start_column=1, end_row=2, end_column=2)
check(s.merged_cell_ranges == ["A1:B2"], "数字指定の結合")
s.unmerge_cells(start_row=1, start_column=1, end_row=2, end_column=2)

# 固定枠: A1 形式(openpyxl と同じ定義)
s.freeze_panes = "B2"
check(s.freeze_panes == "B2", f"固定枠: {s.freeze_panes}")
s.freeze_panes = "A1"  # A1 は「固定なし」
check(s.freeze_panes is None, "A1 で固定が解けない")
s.freeze_panes = "A3"  # 上2行だけ
check(s.freeze_panes == "A3", "行だけの固定")
s.freeze_panes = None

# 改名: 式の参照と名前の定義が追随する(openpyxl は追随しない — うちの上位分)
b2 = office_sheet.Book()
b2[0]["A1"] = 42
w2 = b2.create_sheet("集計")
w2["A1"] = "=Sheet1!A1*2"
check(w2["A1"] == 84, "他のシートへの式")
b2[0].title = "元データ"
check(b2.sheetnames[0] == "元データ", "改名が一覧に出ない")
check(w2.formula("A1") == "=元データ!A1*2", f"改名に式が追随しない: {w2.formula('A1')}")
check(w2["A1"] == 84, "改名後の再計算")
try:
    w2.title = "元データ"
    check(False, "同じ名前への改名が通った")
except ValueError:
    pass
try:
    w2.title = "a[b]"
    check(False, "使えない字のシート名が通った")
except ValueError:
    pass

# 複製・削除・並べ替え・途中に差す
w3 = b2.copy_worksheet(b2[0])
check(w3.title == "元データ Copy", f"複製の名前: {w3.title}")
check(w3["A1"] == 42, "複製に中身が写っていない")
w3["A1"] = 1
check(b2[0]["A1"] == 42, "複製が元とつながったまま(独立していない)")
b2.remove(w3)
check("元データ Copy" not in b2.sheetnames, "削除できていない")
head = b2.create_sheet("先頭", 0)
check(b2.sheetnames[0] == "先頭" and head.title == "先頭", f"途中に差す: {b2.sheetnames}")
b2.move_sheet("先頭", offset=1)
check(b2.sheetnames[1] == "先頭", f"move_sheet の相対のずらし: {b2.sheetnames}")

# 最後の1枚は抜けない(正直に断る)
b3 = office_sheet.Book()
try:
    b3.remove(b3[0])
    check(False, "最後の1枚が抜けた")
except ValueError:
    pass

# openpyxl と読み合う: うちの結合・固定枠が本家に見え、本家の物がうちに見える
if openpyxl is not None:
    with tempfile.TemporaryDirectory() as t:
        out = os.path.join(t, "gokan2.xlsx")
        b4 = office_sheet.Book()
        s4 = b4[0]
        s4["A1"] = "見出し"
        s4.merge_cells("A1:C1")
        s4.freeze_panes = "A2"
        b4.save(out)
        rs = openpyxl.load_workbook(out).active
        check([str(r) for r in rs.merged_cells.ranges] == ["A1:C1"],
              f"うちの結合を openpyxl が読めない: {list(rs.merged_cells.ranges)}")
        check(rs.freeze_panes == "A2", f"うちの固定枠を openpyxl が読めない: {rs.freeze_panes}")

        out2 = os.path.join(t, "opx2.xlsx")
        wb5 = openpyxl.Workbook()
        ws5 = wb5.active
        ws5["A1"] = "題"
        ws5.merge_cells("A1:B2")
        ws5.freeze_panes = "B2"
        wb5.save(out2)
        s5 = office_sheet.Book.open(out2)[0]
        check(s5.merged_cell_ranges == ["A1:B2"], "本家の結合をうちが読めない")
        check(s5.freeze_panes == "B2", f"本家の固定枠をうちが読めない: {s5.freeze_panes}")

# ==================== 第3歩: 画像(xlsx)・表の書き(docx)・polars

# --- add_image: 貼って・見えて・保存で xl/media に入り・読み直しで戻る -----------
PNG_1x1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)
b6 = office_sheet.Book()
s6 = b6[0]
s6["A1"] = "グラフの下じき"
s6.add_image(PNG_1x1, "B2")
check(s6.images == [("B2", 1.0, 1.0)], f"貼った画像が見えない: {s6.images}")
s6.add_image(PNG_1x1, "D4", width_px=200, height_px=100)  # 大きさの上書き
with tempfile.TemporaryDirectory() as t:
    out = os.path.join(t, "img.xlsx")
    b6.save(out)
    s7 = office_sheet.Book.open(out)[0]
    check(("B2", 1.0, 1.0) in s7.images and ("D4", 200.0, 100.0) in s7.images,
          f"画像が往復しない: {s7.images}")
    import zipfile as _zf
    with _zf.ZipFile(out) as z:
        media = [n for n in z.namelist() if n.startswith("xl/media/")]
    check(len(media) == 2, f"xl/media に絵が入っていない: {media}")

    # 径路の文字列でも渡せる(matplotlib の savefig の出口をそのまま)
    p = os.path.join(t, "e.png")
    with open(p, "wb") as f:
        f.write(PNG_1x1)
    s6.add_image(p, "F1")
    check(len(s6.images) == 3, "径路の add_image が効かない")

# --- docx: add_table / add_row / add_column ------------------------------------
d_j = office_doc.Doc()
t_j = d_j.add_table(2, 3)
check(t_j.shape == (2, 3), f"add_table の形: {t_j.shape}")
t_j.cell(0, 0).text = "品名"
row = t_j.add_row()
check(t_j.shape == (3, 3) and len(row.cells) == 3, "add_row")
row.cells[0].text = "ザボガードF"
t_j.add_column()
check(t_j.shape == (3, 4), f"add_column: {t_j.shape}")
try:
    d_j.add_table(2, 2, style="Table Grid")
    check(False, "add_table の style が黙って捨てられた")
except NotImplementedError:
    pass

if pydocx is not None:
    with tempfile.TemporaryDirectory() as t:
        out = os.path.join(t, "hyo.docx")
        d_j.save(out)
        back = pydocx.Document(out)
        bt = back.tables[0]
        check(len(bt.rows) == 3, f"うちが組んだ表を本家が読めない: {len(bt.rows)} 行")
        check(len(bt.rows[0].cells) == 4, f"本家の見た列数: {len(bt.rows[0].cells)}")
        check(bt.cell(0, 0).text == "品名" and bt.cell(2, 0).text == "ザボガードF",
              "うちが書いた中身を本家が読めない")

# --- polars: 第一の変換(ソケットに出ない純関数を直接確かめる)------------------
try:
    import polars as pl
except ImportError:
    pl = None
    print("polars が無いので飛ばした", file=sys.stderr)

if pl is not None:
    grid = [["品名", "数"], ["ザボガードF", 4.0], ["ドリル", 2.0]]
    df = xw._grid_to_frame(grid, pl.DataFrame)
    check(df.columns == ["品名", "数"] and df.shape == (2, 2),
          f"polars への変換: {df.columns} {df.shape}")
    check(df["数"].to_list() == [4.0, 2.0], "polars の中身")
    check(xw._to_grid(df) == grid, f"polars からの往復: {xw._to_grid(df)}")
    check(xw._to_grid(pl.Series([1, 2])) == [[1], [2]], "polars の Series")
    # 見出しなし
    df2 = xw._grid_to_frame([[1, 2], [3, 4]], pl.DataFrame, header=False)
    check(df2.shape == (2, 2), "polars header=False")

try:
    import pandas as pd
except ImportError:
    pd = None
    print("pandas が無いので飛ばした", file=sys.stderr)

if pd is not None:
    grid = [["品名", "数"], ["ザボガードF", 4.0]]
    df3 = xw._grid_to_frame(grid, pd.DataFrame)
    check(list(df3.index) == ["ザボガードF"], "pandas の index(従来どおり)")

print("OK")
