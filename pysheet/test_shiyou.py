# **本家の受け入れ仕様から起こした検査**(2026-08-13)。
#
# python-docx は features/*.feature に「利用者から見た約束」を Gherkin で
# 書いている。その約束をうちの口で確かめる — 出所は NOTICE.md に書いた。
# 写したのは**約束**であって、向こうのコードではない。
#
# ここは test_doc.py(うちの口が動くか)とは役目が違う。**本家の定義に
# 照らして合っているか**だけを見る。合っていない所は台帳に穴として残す
# (docs/pysheet-gokan.ja.md の「本家の定義とずれている所」)。
#
# 手で回すなら:
#   .venv/bin/python pysheet/test_shiyou.py
import sys

from officework import doc as od


def check(cond, msg):
    if not cond:
        print(f"NG: {msg}", file=sys.stderr)
        sys.exit(1)


def raises(exc, f, msg):
    try:
        f()
    except exc:
        return
    except Exception as e:
        check(False, f"{msg}(別の例外: {type(e).__name__}: {e})")
    check(False, f"{msg}(何も起きなかった)")


# ── 段落 ──────────────────────────────────────────────────────
# doc-add-paragraph: 字も様式も省ける。字を渡せばその字を持つ
d = od.Doc()
p0 = d.add_paragraph()
check(p0.text == "", f"空の段落に字がある: {p0.text!r}")
p1 = d.add_paragraph("いろは")
check(p1.text == "いろは", "渡した字が段落に入らない")
check(d.paragraphs[-1].text == "いろは", "足した段落が末尾に来ない")

# doc-add-paragraph: 様式は名前でも様式の物でも渡せる
p2 = d.add_paragraph("見出しにする", style="heading1")
check(p2.style == "heading1", f"名前で渡した様式が付かない: {p2.style}")
p3 = d.add_paragraph("これも", style=d.styles["heading2"])
check(p3.style == "heading2", f"様式の物で渡せない: {p3.style}")

# par-clear-paragraph: 字は消えるが**段落の性質は残る**
pc = d.add_paragraph("消される字", style="heading1")
pc.align = "center"
check(pc.clear() is pc, "clear が自分を返さない(本家と同じ定義)")
check(pc.text == "", "clear で字が消えない")
check(pc.style == "heading1" and pc.align == "center", "clear で段落の性質が飛んだ")

# par-set-text: 字を入れ替えても様式は残る
ps = d.add_paragraph("もとの字", style="heading2")
ps.text = "あとの字"
check(ps.text == "あとの字" and ps.style == "heading2", "字の入れ替えで様式が飛んだ")

# par-insert-paragraph: 前に差し込むと1つ増え、その位置に入る
n = len(d.paragraphs)
pi = ps.insert_paragraph_before("差し込んだ", style="heading3")
check(len(d.paragraphs) == n + 1, "差し込んでも段落が増えない")
i = [x.text for x in d.paragraphs].index("差し込んだ")
check(d.paragraphs[i + 1].text == "あとの字", "差し込みが指した段落の前に入っていない")
check(pi.style == "heading3", "差し込んだ段落の様式が付かない")

# doc-add-heading: 既定は1段目。level を渡せばその段
h1 = d.add_heading("見出し")
check(h1.style == "heading1", f"add_heading の既定が1段目でない: {h1.style}")
check(h1.text == "見出し", "見出しの字が入らない")
check(d.add_heading("2段目", 2).style == "heading2", "level=2 が2段目でない")
check(d.add_heading("3段目", 3).style == "heading3", "level=3 が3段目でない")
# **うちの見出しは3段まで**(模型の粒度)。本家の 0=Title・4〜9 は正直に断る
raises(ValueError, lambda: d.add_heading("題", 0), "level=0 を黙って受けている")
raises(ValueError, lambda: d.add_heading("題", 4), "level=4 を黙って受けている")
raises(ValueError, lambda: d.add_heading("題", 10), "範囲の外を黙って受けている")

# doc-add-page-break: 改ページだけの段落が末尾に付く
before = len(d.paragraphs)
d.add_page_break()
check(len(d.paragraphs) == before + 1, "改ページで段落が増えない")
check(d.paragraphs[-1].paragraph_format.page_break_before is True,
      "改ページの段落に page_break_before が立っていない")

# ── run ───────────────────────────────────────────────────────
# par-add-run: 字を渡せばその字を持つ run が末尾に付く
pr = d.add_paragraph("あ")
r = pr.add_run("い")
check(r.text == "い", "add_run に渡した字が入らない")
check(pr.text == "あい", "run を足しても段落の字に出てこない")

# run の書式は**全部の口が書ける**(台帳の約束)。size_pt は包み(_doc.py)の
# setter だけが漏れていて、Rust 側にはあるのに Python から書けなかった
# (2026-08-14 に発見)— 口ごとに縛って、片方だけの漏れを捕まえる
r.size_pt = 14
check(r.size_pt == 14.0, f"size_pt が書けない: {r.size_pt}")
r.size_pt = None
check(r.size_pt is None, "size_pt の指定を外せない(None = 文書の既定)")

# run-clear-run: 字は消えるが**書式は残る**
r.bold = True
r.italic = True
check(r.clear() is r, "run.clear が自分を返さない")
check(r.text == "", "run.clear で字が消えない")
check(r.bold is True and r.italic is True, "run.clear で書式が飛んだ")

# run-add-content / txt-add-break: 改行・タブは run の中身として順に並ぶ
rb = d.add_paragraph("").add_run("あ")
rb.add_break()
rb.add_text("い")
rb.add_tab()
rb.add_text("う")
kinds = [type(x).__name__ if not isinstance(x, str) else "str"
         for x in rb.iter_inner_content()]
check(kinds == ["str", "Break", "str", "Tab", "str"],
      f"run の中身が順に返らない: {kinds}")

# run-char-style: 文字スタイルは名前でも様式の物でも
d.styles.add_style("強い", "character")
rs = d.add_paragraph("").add_run("字")
rs.style = "強い"
check(rs.style == "強い", f"文字スタイルが付かない: {rs.style}")
# 無い名前は**黙って作らない**(add_style で作ってから、が家の作法)
raises(Exception, lambda: setattr(rs, "style", "無い様式"), "無い様式を黙って受けている")

# ── 表 ────────────────────────────────────────────────────────
# doc-add-table: 行数・列数のとおりに出来る
t = d.add_table(rows=2, cols=3)
check(t.shape == (2, 3), f"表の形が違う: {t.shape}")
check(len(t.rows) == 2 and len(t.columns) == 3, "行・列の数が合わない")

# tbl-cell-text / tbl-cell-access: セルの字は代入でき、行・列から引ける
t.cell(0, 0).text = "左上"
t.cell(1, 2).text = "右下"
check(t.cell(0, 0).text == "左上", "セルに入れた字が読めない")
check(t.rows[0].cells[0].text == "左上", "行から引いたセルが食い違う")
check(t.columns[2].cells[1].text == "右下", "列から引いたセルが食い違う")
check(t[1][2].text == "右下", "添字で引いたセルが食い違う")

# tbl-add-row-or-col: 足した行は同じ列数を持つ
t.add_row()
check(t.shape == (3, 3), f"行を足しても形が変わらない: {t.shape}")
check(len(t.rows[-1].cells) == 3, "足した行の列数が違う")
t.add_column()
check(t.shape == (3, 4), f"列を足しても形が変わらない: {t.shape}")
check(len(t.columns[-1].cells) == 3, "足した列の行数が違う")
# 等分の表に幅つきの列を足すときは、**今ある列を等分で埋めてから**足します
# (2026-08-28。前は断っていましたが、本家では通る書き方でした)。
# 埋める値は紙の幅から余白を引いた物を列の数で割った物で、見た目は
# 変わりません
mae = t.shape[1]
t.add_column(od.Mm(30))
check(t.shape[1] == mae + 1, f"幅つきの列が足せない: {t.shape}")
check(abs(t.cell(0, mae).width.mm - 30) < 0.1,
      f"足した列の幅が入らない: {t.cell(0, mae).width}")
check(t.cell(0, 0).width is not None, "今ある列の幅が埋まっていない")


# tbl-props: autofit は明示が無ければ True(本家の「no explicit → autofit」)
check(t.autofit is True, f"autofit の既定が True でない: {t.autofit}")
t.autofit = False
check(t.autofit is False, "autofit を False にできない")

# ── 断ると決めた所(黙って落とさない)──────────────────────────
# cmt-* : 本家は run から run までを範囲にするが、うちのコメントは段落単位
rm = d.add_paragraph("字").add_run("あ")
raises(NotImplementedError, lambda: rm.mark_comment_range(rm, 1),
       "範囲コメントを黙って受けている")
# txt-parfmt-props: 段落の前後の余白は 2026-08-27 に模型へ入りました
# (台帳 #5)。断るのをやめて、往復するかを見ます
pf = d.add_paragraph("字").paragraph_format
pf.space_before = od.Pt(12)
pf.space_after = 6          # 生の数(pt)でも受ける
check(pf.space_before.pt == 12.0, f"段落前の余白が入らない: {pf.space_before}")
check(pf.space_after.pt == 6.0, f"段落後の余白が入らない: {pf.space_after}")
# 字下げも 2026-08-27 に入りました。**左は模型では段数**(1段=全角2字)
# なので、本文の既定の大きさ(10.5pt)で数えていちばん近い段に寄ります。
# 1行目の字下げは docx と同じ長さのまま往復します
pf.first_line_indent = od.Pt(10.5)
check(abs(pf.first_line_indent.pt - 10.5) < 0.1,
      f"1行目の字下げが入らない: {pf.first_line_indent}")
pf.left_indent = od.Pt(21)          # 全角2字ぶん = 1段
check(abs(pf.left_indent.pt - 21) < 0.1, f"左の字下げが入らない: {pf.left_indent}")
pf.left_indent = od.Pt(42)          # 2段
check(abs(pf.left_indent.pt - 42) < 0.1, f"2段目が入らない: {pf.left_indent}")
pf.left_indent = None
check(pf.left_indent.pt == 0.0, f"字下げが外れない: {pf.left_indent}")

# ── 無指定は無指定のまま往復する(2026-08-13 に塞いだ穴)──────
# 指定の無い文字の大きさが往復で 10.5pt に焼き付いていた(w:sz を必ず
# 書く + Run.size_pt が f32 で無指定を持てない)。Run.size_pt を
# Option にして根治 — 本家の font.size が None を返すのと同じ約束。
try:
    import docx as _honke
except ImportError:
    _honke = None

if _honke is not None:
    import os
    import tempfile

    with tempfile.TemporaryDirectory() as t:
        moto = os.path.join(t, "moto.docx")
        ato = os.path.join(t, "ato.docx")
        hd = _honke.Document()
        hd.add_paragraph("大きさを指定していない字")
        hd.save(moto)
        check(_honke.Document(moto).paragraphs[0].runs[0].font.size is None,
              "本家が作った時点で大きさが入っている(前提が崩れた)")
        od.Doc.open(moto).save(ato)
        ima = _honke.Document(ato).paragraphs[0].runs[0].font.size
        check(ima is None,
              f"指定の無い大きさが往復で焼き付いた(本家の読み: {ima!r}。"
              "None のままが正 — 2026-08-13 に塞いだ穴が開き直している)")
        # うちの口でも同じ約束(size_pt は None = 指定なし)
        check(od.Doc.open(moto).paragraphs[0].runs[0].size_pt is None,
              "無指定の run の size_pt が None でない")

print("OK")


def 文字は文字のまま置かれる():
    """**ファイルの口と打鍵の口は別物**(2026-08-15)。

    前は Python から文字を置くと「人が打った字の解釈器」を通っていて、
    `"0001"` が数の 1 になり、前後の空白も削られていた。品番・郵便番号・
    電話番号・会員番号が壊れる所。openpyxl は文字を文字のまま置く。
    """
    from officework import sheet

    b = sheet.Book()
    ws = b.active
    ws["A1"] = "0001"
    ws["A2"] = " 山田 太郎 "
    ws["A3"] = "TRUE"
    ws["A4"] = "3.14"
    ws["A5"] = "=1+2"
    ws["A6"] = ""
    assert ws["A1"].value == "0001", f"品番が数にされた: {ws['A1'].value!r}"
    assert ws["A2"].value == " 山田 太郎 ", f"前後の空白が削られた: {ws['A2'].value!r}"
    assert ws["A3"].value == "TRUE", f"TRUE が真偽にされた: {ws['A3'].value!r}"
    assert ws["A4"].value == "3.14", f"数に見える字が数にされた: {ws['A4'].value!r}"
    # = で始まる字だけは式(openpyxl も同じ)
    assert ws.formula("A5") == "=1+2", f"式にならない: {ws.formula('A5')!r}"
    assert ws["A5"].value == 3, f"式が計算されない: {ws['A5'].value!r}"
    # 空文字は空のセル(使っている範囲の数え方を変えないため)
    assert ws["A6"].value is None, f"空文字が値になった: {ws['A6'].value!r}"
    print("  文字は文字のまま置かれる: ok")


def 整数はintで返る():
    """openpyxl は xlsx の <v>340</v> を int で返す。340.0 だと見せる前に
    毎回 int() が要り、品番や個数の桁が汚れる(2026-08-15)"""
    from officework import sheet

    b = sheet.Book()
    ws = b.active
    ws["A1"] = 340
    ws["A2"] = 3.5
    ws["A3"] = -7
    assert isinstance(ws["A1"].value, int) and ws["A1"].value == 340, repr(ws["A1"].value)
    assert isinstance(ws["A2"].value, float) and ws["A2"].value == 3.5, repr(ws["A2"].value)
    assert isinstance(ws["A3"].value, int) and ws["A3"].value == -7, repr(ws["A3"].value)
    print("  整数はintで返る: ok")


def appendは空の行も進める():
    """`append([])` で行が進まないと、表題・空行・記入欄の定型の用紙が
    1行ずつずれる(2026-08-15。種苗の会の注文書の見本で踏んだ)"""
    from officework import sheet

    b = sheet.Book()
    ws = b.active
    ws.append(["あ"])
    ws.append([])
    ws.append([""])
    ws.append(["い"])
    assert ws.cell(row=1, column=1).value == "あ", "1行目"
    assert ws.cell(row=2, column=1).value is None, "空の行が飛んだ"
    assert ws.cell(row=4, column=1).value == "い", f"4行目にならない"
    print("  appendは空の行も進める: ok")


def 範囲の参照は組の組で返る():
    """`for row in ws["A1:C1"]` は openpyxl の定番の書き方(2026-08-15)"""
    from officework import sheet

    b = sheet.Book()
    ws = b.active
    ws.append(["a", "b", "c"])
    ws.append(["d", "e", "f"])
    r = ws["A1:C1"]
    assert len(r) == 1 and len(r[0]) == 3, f"形が違う: {r}"
    assert [c.value for c in r[0]] == ["a", "b", "c"]
    col = ws["A1:A2"]
    assert len(col) == 2 and len(col[0]) == 1, f"1列の形が違う: {col}"
    assert [row[0].value for row in col] == ["a", "d"]
    print("  範囲の参照は組の組で返る: ok")


def 径路はPathでも受ける():
    """openpyxl は pathlib.Path を受ける(2026-08-15)"""
    import pathlib
    import tempfile

    from officework import sheet

    with tempfile.TemporaryDirectory() as d:
        p = pathlib.Path(d) / "t.xlsx"
        b = sheet.Book()
        b.active["A1"] = "あ"
        b.save(p)
        assert p.exists(), "Path で保存できない"
        b2 = sheet.Book.open(p)
        assert b2[b2.sheet_names[0]]["A1"].value == "あ", "Path で開けない"
    print("  径路はPathでも受ける: ok")


文字は文字のまま置かれる()
整数はintで返る()
appendは空の行も進める()
範囲の参照は組の組で返る()
径路はPathでも受ける()
