#!/usr/bin/env python3
"""**出発点の3つに対して、いま何ができているか。**

    .venv/bin/python tools/cover_check.py

発注者(2026-08-29)「出発点を Euro-Office と python-docx と openpyxl に
している。これらができることを、すべてできるようにするのが今回の目的」。

100%の互換を目指すと複雑になりすぎるので、**この3つができること**を
完了の定義にします。そこまで来ているかを数える道具です。

openpyxl と python-docx は**同じ名前の物があるか**を見ます。名前だけの
一致なので、中身が正しいかは別の試験(test_gokan.py など)の受け持ちです。
ここは「まだ触っていない口」を見つけるための粗い網です。

Euro-Office のリボンは Rust の側で数えます:

    cargo run -q -p face --example cover
"""
import inspect
import sys

from officework import _doc as od
from officework import sheet as ow

# 見ない名前。**本家の内部の作りに由来する物**で、こちらが持つ意味がない
NOZOKU = {
    # openpyxl の内部
    "path", "parser", "write_only", "encoding", "iso_dates", "epoch",
    "mime_type", "template", "excel_base_date",
    # python-docx の内部
    "element", "part", "_p", "_r", "_tbl",
}


def cls_members(cls):
    return {n for n, _ in inspect.getmembers(cls) if not n.startswith("_")}


def obj_members(obj):
    return {n for n in dir(obj) if not n.startswith("_")}


def kurabe(midashi, honke_cls, uchi_obj):
    honke = cls_members(honke_cls) - NOZOKU
    uchi = obj_members(uchi_obj)
    nai = sorted(honke - uchi)
    aru = len(honke) - len(nai)
    wari = aru * 100 // max(1, len(honke))
    print(f"{midashi:34} {aru:3} / {len(honke):3}  ({wari}%)")
    return nai


def chart_kurabe():
    """**チャートは種類で数えます。**

    本家(openpyxl)は OOXML のチャート XML の模型を持ち、こちらは図形の
    集まりとして描きます。作りが違うので、欄の名前を突き合わせても意味が
    ありません(本家の欄のほとんどは Excel に描かせるための指図です)。

    利用者から見て同じなのは**どの種類のグラフが出せるか**なので、そこを
    数えます。
    """
    import inspect

    import openpyxl.chart as oc

    honke = sorted(
        n for n in dir(oc) if n.endswith("Chart") and inspect.isclass(getattr(oc, n))
    )
    # 本家の名前と、こちらの `add_chart` の `kind` の対応
    taiou = {
        "AreaChart": "area",
        "BarChart": "bar",
        "BubbleChart": "bubble",
        "DoughnutChart": "doughnut",
        "LineChart": "line",
        "PieChart": "pie",
        "ProjectedPieChart": None,
        "RadarChart": "radar",
        "ScatterChart": "scatter",
        "StockChart": None,
        "SurfaceChart": None,
    }
    b = ow.Book()
    ws = b[0]
    for r, v in enumerate([3, 1, 4, 1, 5], start=1):
        ws.cell(r, 1).value = f"項目{r}"
        ws.cell(r, 2).value = v
    aru, nai = [], []
    for na in honke:
        kind = taiou.get(na)
        if kind is None:
            nai.append(na)
            continue
        try:
            ws.add_chart(kind, data="B1:B5", categories="A1:A5", at="D1")
            aru.append(na)
        except Exception:
            nai.append(na)
    wari = len(aru) * 100 // max(1, len(honke))
    print(f"{'openpyxl チャートの種類':34} {len(aru):3} / {len(honke):3}  ({wari}%)")
    return nai


def main():
    ana = {}

    # --- openpyxl ---
    from openpyxl.cell.cell import Cell as OCell
    from openpyxl.workbook import Workbook as OBook
    from openpyxl.worksheet.worksheet import Worksheet as OSheet

    b = ow.Book()
    ana["openpyxl.Workbook"] = kurabe("openpyxl Workbook", OBook, b)
    ana["openpyxl.Worksheet"] = kurabe("openpyxl Worksheet", OSheet, b[0])
    ana["openpyxl.Cell"] = kurabe("openpyxl Cell", OCell, b[0]["A1"])

    # --- python-docx ---
    from docx.document import Document as DDoc
    from docx.table import Table as DTable
    from docx.text.paragraph import Paragraph as DPara
    from docx.text.run import Run as DRun

    d = od.Doc()
    p = d.add_paragraph("x")
    r = p.add_run("y")
    t = d.add_table(1, 1)
    ana["python-docx.Document"] = kurabe("python-docx Document", DDoc, d)
    ana["python-docx.Paragraph"] = kurabe("python-docx Paragraph", DPara, p)
    ana["python-docx.Run"] = kurabe("python-docx Run", DRun, r)
    ana["python-docx.Table"] = kurabe("python-docx Table", DTable, t)

    # --- チャート(種類で数えます)---
    ana["openpyxl.Chart"] = chart_kurabe()

    nokori = sum(len(v) for v in ana.values())
    print(f"\nまだ無い口: {nokori}")
    for k, v in ana.items():
        if v:
            print(f"  {k}")
            for n in v:
                print(f"      {n}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
