# 注文書付きカタログ(カタログ.docx)を商品マスタから作る。中身はすべて架空。
#
#   .venv/bin/python sample/gen_catalog.py                # サーバーから取る
#   .venv/bin/python sample/gen_catalog.py --offline      # 同梱データで作る
#
# 商品マスタの正本はサーバー(catalog_server.py)にあり、docx は生成物 —
# 「見ながら整える仕事は writer、データを作る仕事は Python」の分業の見本。
# サーバーに繋がらなければ、そう言ってから同梱の見本データで作る。
# writer で開いたら 参考資料 > 目次 で、見出しからページ番号つきの目次が作れる。
import csv
import io
import pathlib
import sys
import urllib.request

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 書き出す先は、この .py の隣です(どこで走らせても同じ)
ここ = pathlib.Path(__file__).resolve().parent

URL = "http://127.0.0.1:8765/catalog.csv"

# 同梱の見本データ(サーバーが無い機械でも同じ物を作り直せるように)
FALLBACK = [
    ("A-101", "筆記具", "ボールペン(黒)", "0.7mm・油性", 150),
    ("A-102", "筆記具", "ボールペン(赤)", "0.7mm・油性", 150),
    ("A-103", "筆記具", "ボールペン(青)", "0.7mm・油性", 150),
    ("A-104", "筆記具", "シャープペン", "0.5mm", 220),
    ("A-105", "筆記具", "シャープ替芯", "0.5mm・40本", 120),
    ("A-106", "筆記具", "蛍光マーカー(黄)", "太細両用", 130),
    ("A-107", "筆記具", "蛍光マーカー(桃)", "太細両用", 130),
    ("A-108", "筆記具", "油性ペン(黒)", "太字", 160),
    ("A-109", "筆記具", "鉛筆HB", "12本入り", 480),
    ("A-110", "筆記具", "消しゴム", "まとまるタイプ", 90),
    ("B-201", "紙製品", "コピー用紙A4", "500枚", 550),
    ("B-202", "紙製品", "コピー用紙B5", "500枚", 520),
    ("B-203", "紙製品", "ノートA罫", "30枚・セミB5", 180),
    ("B-204", "紙製品", "レポート用紙A4", "50枚", 250),
    ("B-205", "紙製品", "付箋 75×75mm", "桃・100枚", 210),
    ("B-206", "紙製品", "付箋 75×25mm", "3色・各100枚", 260),
    ("B-207", "紙製品", "封筒 長形3号", "100枚", 680),
    ("B-208", "紙製品", "クラフト封筒 角形2号", "50枚", 750),
    ("C-301", "ファイル・収納", "クリアファイルA4", "10枚", 240),
    ("C-302", "ファイル・収納", "パイプ式ファイルA4", "背幅5cm", 780),
    ("C-303", "ファイル・収納", "個別フォルダA4", "10枚", 620),
    ("C-304", "ファイル・収納", "2穴バインダーA4", "背幅3cm", 450),
    ("C-305", "ファイル・収納", "書類トレーA4", "積み重ね可", 520),
    ("C-306", "ファイル・収納", "マグネットバー", "20cm", 330),
    ("D-401", "事務機器", "電卓", "12桁", 1480),
    ("D-402", "事務機器", "ホッチキス10号", "20枚とじ", 620),
    ("D-403", "事務機器", "ホッチキス針10号", "1000本", 110),
    ("D-404", "事務機器", "2穴パンチ", "20枚", 830),
    ("D-405", "事務機器", "テープカッター", "大巻用", 690),
    ("D-406", "事務機器", "はさみ", "175mm", 420),
    ("D-407", "事務機器", "カッターL型", "替刃1枚付き", 380),
    ("D-408", "事務機器", "スティックのり", "約10g", 140),
    ("E-501", "梱包・雑貨", "ガムテープ(布)", "50mm×25m", 280),
    ("E-502", "梱包・雑貨", "OPPテープ(透明)", "48mm×100m", 190),
    ("E-503", "梱包・雑貨", "緩衝材", "ぷちぷち・10m", 640),
    ("E-504", "梱包・雑貨", "宅配袋(大)", "10枚", 520),
]


def fetch():
    """商品マスタを取る。返りは (行の一覧, 出所の説明)。"""
    if "--offline" not in sys.argv:
        try:
            raw = urllib.request.urlopen(URL, timeout=3).read()
            rows = list(csv.reader(io.StringIO(raw.decode("utf-8"))))[1:]
            rows = [(a, b, c, d, int(e)) for a, b, c, d, e in rows]
            return rows, f"サーバー {URL}"
        except OSError as e:
            print(f"サーバーに繋がりません({e})— 同梱の見本データで作ります")
    return list(FALLBACK), "同梱の見本データ"


def main():
    products, source = fetch()

    d = docx.Document()
    d.add_heading("事務用品カタログ(2026年秋)", level=1)
    d.add_paragraph("例示文具株式会社 — 価格はすべて税抜。ご注文は巻末の注文書で。")
    d.add_paragraph(f"この版の出所: {source}({len(products)}品目)")

    # 分類ごとに見出し+表(出てきた順を保つ)
    categories = []
    for p in products:
        if p[1] not in categories:
            categories.append(p[1])
    for cat in categories:
        items = [p for p in products if p[1] == cat]
        d.add_heading(cat, level=2)
        t = d.add_table(rows=1 + len(items), cols=4)
        for j, h in enumerate(["品番", "品名", "説明", "単価(税抜)"]):
            t.rows[0].cells[j].text = h
        for i, (code, _, name, desc, price) in enumerate(items):
            row = t.rows[1 + i].cells
            row[0].text = code
            row[1].text = name
            row[2].text = desc
            row[3].text = f"{price:,}円"

    # 巻末: 注文書(改ページして1枚に)
    d.add_page_break()
    d.add_heading("注文書", level=1)
    d.add_paragraph("例示文具株式会社 行(FAX 012-345-0000)")

    t = d.add_table(rows=2, cols=4)
    for j, (k, v) in enumerate([("社名", ""), ("担当", "")]):
        t.rows[0].cells[j * 2].text = k
        t.rows[0].cells[j * 2 + 1].text = v
    for j, (k, v) in enumerate([("電話", ""), ("納品希望日", "")]):
        t.rows[1].cells[j * 2].text = k
        t.rows[1].cells[j * 2 + 1].text = v

    d.add_paragraph("下の表に品番と数量をご記入ください(品番はカタログ各ページの表から)。")
    t = d.add_table(rows=11, cols=4)
    for j, h in enumerate(["品番", "品名", "数量", "金額(税抜)"]):
        t.rows[0].cells[j].text = h

    d.add_paragraph("合計(税込):            円")
    p = d.add_paragraph("以上")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    d.save(ここ / "カタログ.docx")
    print(f"書いた: sample/カタログ.docx({source}・{len(products)}品目・{len(categories)}分類)")


if __name__ == "__main__":
    main()
