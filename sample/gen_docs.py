# 公開できるサンプル文書(docx)をまとめて作る。中身はすべて架空。
#
#   .venv/bin/python sample/gen_docs.py
#
# docx は python-docx で作る(SEKKEI「writer には橋を作らない」—
# その保存を writer が読めることは実物で確認済み)。
# サンプルは生成物 — 直すのはこのファイル。
import pathlib

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 書き出す先は、この .py の隣です(どこで走らせても同じ)
ここ = pathlib.Path(__file__).resolve().parent


def report():
    d = docx.Document()
    d.add_heading("月次業務報告(2026年7月)", level=1)

    d.add_heading("概況", level=2)
    d.add_paragraph(
        "7月の受注は3件(見積提出5件)。外壁塗装の引き合いが続いており、"
        "8月は足場資材の手配が山になる見込み。"
    )

    d.add_heading("受注の実績", level=2)
    t = d.add_table(rows=4, cols=4)
    for j, h in enumerate(["受注日", "件名", "発注元", "金額(税込)"]):
        t.rows[0].cells[j].text = h
    for i, row in enumerate([
        ("7月3日", "外壁塗装工事", "株式会社みほん商事", "640,200円"),
        ("7月15日", "屋根の補修", "例示ビル管理", "231,000円"),
        ("7月28日", "駐車場の白線引き直し", "架空団地自治会", "88,000円"),
    ]):
        for j, v in enumerate(row):
            t.rows[i + 1].cells[j].text = v

    d.add_heading("来月の予定", level=2)
    d.add_paragraph("・8月上旬: みほん商事の現場に着手(足場は第1週に設置)")
    d.add_paragraph("・8月中旬: 夏季休業(11日〜15日)")
    d.add_paragraph("・8月下旬: 見積フォローの巡回")

    p = d.add_paragraph()
    p.add_run("以上").bold = True
    d.save(ここ / "報告書.docx")
    print("書いた: sample/報告書.docx")


def cover_letter():
    """送付状 — 揃え(右・中央)と定型の構えの見本。"""
    d = docx.Document()
    right = d.add_paragraph("2026年8月4日")
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    d.add_paragraph("株式会社みほん商事\n総務部 御中")

    sender = d.add_paragraph("例示工務店\n見本県架空市例示町1-2-3\n電話 012-345-6789")
    sender.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title = d.add_paragraph()
    r = title.add_run("書類送付のご案内")
    r.bold = True
    r.font.size = docx.shared.Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.add_paragraph(
        "拝啓 時下ますますご清栄のこととお慶び申し上げます。"
        "平素は格別のお引き立てを賜り、厚く御礼申し上げます。"
    )
    d.add_paragraph("さて、下記の書類をお送りいたしますので、ご査収のほどお願い申し上げます。")
    p = d.add_paragraph("敬具")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ki = d.add_paragraph()
    ki.add_run("記").bold = True
    ki.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("・御見積書(M-2026-001)　1通")
    d.add_paragraph("・会社案内　1部")
    p = d.add_paragraph("以上")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    d.save(ここ / "送付状.docx")
    print("書いた: sample/送付状.docx")


def minutes():
    """議事録 — 見出しの階層と表、決定事項の見本。"""
    d = docx.Document()
    d.add_heading("定例会議 議事録(2026年8月4日)", level=1)

    t = d.add_table(rows=3, cols=2)
    for i, (k, v) in enumerate([
        ("日時", "2026年8月4日(火)10:00〜10:45"),
        ("場所", "事務所 打ち合わせ卓"),
        ("出席", "見本太郎・例示花子・架空次郎"),
    ]):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v

    d.add_heading("議題1: みほん商事の現場の段取り", level=2)
    d.add_paragraph(
        "足場は8月第1週に設置(例示塗装へ発注済み)。近隣への挨拶回りは"
        "前週の金曜までに済ませる。担当は例示花子。"
    )

    d.add_heading("議題2: 夏季休業の告知", level=2)
    d.add_paragraph("8月11日〜15日を休業とし、留守番電話と貼り紙で告知する。")

    d.add_heading("決定事項", level=2)
    d.add_paragraph("・挨拶回り: 8月1日(金)まで(例示花子)")
    d.add_paragraph("・休業の貼り紙: 8月6日(水)まで(架空次郎)")
    d.add_paragraph("・次回定例: 8月18日(月)10:00")

    d.save(ここ / "議事録.docx")
    print("書いた: sample/議事録.docx")


if __name__ == "__main__":
    report()
    cover_letter()
    minutes()
