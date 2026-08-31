# e-shop プラグインの試作 — 「販売者が Word で店を書く」の一本道。
#
#   .venv/bin/python sample/eshop.py [カタログ.docx]
#
# 文書(カタログ)から商品マスタを取り出して sample/商品マスタ.csv に書く。
# catalog_server.py がそれを読んで店(JS なしの Web)を開く:
#
#   .venv/bin/python sample/eshop.py          # 文書 → 商品マスタ.csv
#   python3 sample/catalog_server.py          # 店を開く(CSV を自動で読む)
#
# 店の正本は文書 — 品を足す・値を直すのは writer(または Word)で文書を
# 直し、この2つを回し直すだけ。Web の知識は要らない。
# 本実装は writer のプラグイン(~/.config/officework/plugins・サンドボックスつき)になる —
# そのとき d(python-docx の文書)束縛でこの中身がそのまま使える。
import csv
import pathlib
import re
import sys

import docx

ここ = pathlib.Path(__file__).resolve().parent
SRC = sys.argv[1] if len(sys.argv) > 1 else ここ / "カタログ.docx"
OUT = ここ / "商品マスタ.csv"


def main():
    d = docx.Document(SRC)

    # 本文の並び順に「見出し2 → 表」を対応づける(分類は直前の見出し)。
    # 商品表の見出し行は 品番/品名/説明/単価 — それ以外の表(記入欄など)は飛ばす。
    headings = iter(p.text for p in d.paragraphs if p.style.name == "Heading 2")
    products, skipped = [], 0
    cat = ""
    body = d.element.body
    tbl_i = 0
    for child in body:
        if child.tag.endswith("}p"):
            text = "".join(n.text or "" for n in child.iter() if n.tag.endswith("}t"))
            for h in [p for p in d.paragraphs if p._p is child]:
                if h.style.name == "Heading 2":
                    cat = text
        elif child.tag.endswith("}tbl"):
            t = d.tables[tbl_i]
            tbl_i += 1
            head = [c.text.strip() for c in t.rows[0].cells]
            if head[:2] != ["品番", "品名"]:
                skipped += 1
                continue
            for row in t.rows[1:]:
                code, name, desc, price = (c.text.strip() for c in row.cells[:4])
                m = re.search(r"[\d,]+", price)
                if code and m:
                    products.append(
                        (code, cat or "その他", name, desc, int(m.group().replace(",", ""))))

    if not products:
        print(f"商品表が見つかりません({SRC} — 見出し行が 品番/品名/… の表が要る)")
        sys.exit(1)

    with open(OUT, "w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["品番", "分類", "品名", "説明", "単価"])
        w.writerows(products)
    cats = []
    for p in products:
        if p[1] not in cats:
            cats.append(p[1])
    print(f"書いた: {OUT}({len(products)} 品目・{len(cats)} 分類。"
          f"商品表でない表 {skipped} 件は飛ばした)")
    print("次: python3 sample/catalog_server.py で店が開く(この CSV を自動で読む)")


if __name__ == "__main__":
    main()
