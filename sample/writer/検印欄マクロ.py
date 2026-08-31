# writer のマクロの見本 — 文書の末尾に検印欄(承認・確認・担当)を足す。
#
# 使い方: writer で文書を開き、プラグイン > マクロ でこの .py を選ぶ。
# 台本には d(python-docx の Document)が束縛されていて、d への変更が
# 1手として文書に入る(Ctrl+Z で戻る)。保存はしない — writer 側の仕事。
# 実行はサンドボックス(bubblewrap)の中。コードは文書には載らない(docx に実行
# コードを入れない — 詳しくは docs/en/python-manual.adoc)。
#
# 単体でも試せる: .venv/bin/python sample/writer/検印欄マクロ.py 文書.docx


def 検印欄を足す(doc):
    doc.add_paragraph("")
    t = doc.add_table(rows=2, cols=3)
    for j, h in enumerate(["承認", "確認", "担当"]):
        t.rows[0].cells[j].text = h
        t.rows[1].cells[j].text = "　"  # 印を押す空き(全角空白で高さを確保)


if "d" in globals():                     # writer のマクロとして
    検印欄を足す(d)
    print("検印欄を末尾に足しました(Ctrl+Z で戻せます)")
else:                                    # 単体で(検証にも使う)
    import pathlib
    import sys

    import docx

    src = sys.argv[1] if len(sys.argv) > 1 else str(
        pathlib.Path(__file__).resolve().parent.parent / "報告書.docx")
    doc = docx.Document(src)
    検印欄を足す(doc)
    out = src.replace(".docx", "_検印.docx")
    doc.save(out)
    print("書いた:", out)
