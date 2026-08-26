#!/usr/bin/env python3
"""同じコードで docx を2つ作る — python-docx と officework.doc の差し替え試験。

    .venv/bin/python test/same_code_docx.py [出力先(既定 test/out)]

build() は1つだけ。渡すライブラリの束(Document・Mm・揃えの値)を
替えて2回呼び、出来た2つの docx を python-docx の目で読み比べます。
文面が同じなら「import の1行を替えれば同じコードが動く」の証明になります。
"""
import sys
from pathlib import Path

OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).resolve().parent / "out"
OUT.mkdir(parents=True, exist_ok=True)


def build(lib, out):
    """通知文を1枚作る。ここは2つのライブラリで完全に同じコード。"""
    d = lib["Document"]()
    s = d.sections[0]
    s.page_width, s.page_height = lib["Mm"](210), lib["Mm"](297)
    s.top_margin = s.bottom_margin = lib["Mm"](20)
    s.left_margin = s.right_margin = lib["Mm"](20)

    p = d.add_paragraph("2026年8月26日")
    p.alignment = lib["RIGHT"]
    d.add_paragraph("会員各位")

    h = d.add_heading("定例会 開催のお知らせ", level=1)
    h.alignment = lib["CENTER"]

    p = d.add_paragraph("下記のとおり定例会を開きます。")
    r = p.add_run("出欠のご返事は 9月5日(金)までにお願いします。")
    r.bold = True
    r.underline = True

    p = d.add_paragraph("記")
    p.alignment = lib["CENTER"]

    t = d.add_table(rows=3, cols=2)
    rows = [("日時", "9月12日(土) 14:00〜16:00"),
            ("場所", "工業技術センター 2階 会議室"),
            ("議題", "新しい道具の紹介と実演")]
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v

    p = d.add_paragraph("以上")
    p.alignment = lib["RIGHT"]

    d.core_properties.title = "定例会の開催通知"
    d.save(str(out))


# --- python-docx の束 ---
import docx as pydocx
from docx.shared import Mm as PdMm
from docx.enum.text import WD_ALIGN_PARAGRAPH

lib_pydocx = {"Document": pydocx.Document, "Mm": PdMm,
              "CENTER": WD_ALIGN_PARAGRAPH.CENTER, "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT}

# --- officework の束 ---
from officework import doc as owdoc

lib_ow = {"Document": owdoc.Doc, "Mm": owdoc.Mm,
          "CENTER": "center", "RIGHT": "right"}

a = OUT / "same_code_pydocx.docx"
b = OUT / "same_code_officework.docx"
build(lib_pydocx, a)
build(lib_ow, b)
print("2枚書けた:", a.name, "/", b.name)

# --- python-docx の目で読み比べる ---
def snapshot(path):
    w = pydocx.Document(str(path))
    paras = [(p.text, p.style.name, str(p.alignment)) for p in w.paragraphs]
    cells = [c.text for t in w.tables for row in t.rows for c in row.cells]
    return paras, cells

pa, ca = snapshot(a)
pb, cb = snapshot(b)
bad = 0
if ca != cb:
    bad += 1
    print("× 表が違う:", ca, "≠", cb)
if len(pa) != len(pb):
    bad += 1
    print(f"× 段落の数が違う: {len(pa)} ≠ {len(pb)}")
for x, y in zip(pa, pb):
    if x != y:
        bad += 1
        print(f"× 段落が違う: {x} ≠ {y}")
if bad:
    raise SystemExit(f"読み比べ: 違い {bad} 件")
print("読み比べ: 段落・様式・揃え・表とも同じ(python-docx の目)")
