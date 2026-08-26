#!/usr/bin/env python3
"""python-docx の「書く」機能を、実在しそうな1ページの文書5種類に詰めて作る。

    .venv/bin/python test/write_docx_pydocx.py [出力先フォルダ(既定 test/out)]

出来上がり(それぞれ1枚):
  開催通知.docx   スタイル・見出し・字の書式(太字/下線/色/大きさ)・
                  右揃え・字下げ・タブ・行間・A4 と余白
  議事録.docx     表(スタイル・結合・列幅・行の高さ)・箇条書きと番号・
                  中央揃えの表・セルの縦の揃え
  操作手順書.docx 番号つきの手順・字の等幅スタイル・画像・改ページ・
                  2ページ目の頭に「ここから別紙」
  送付状.docx     ヘッダーとフッター・1枚目だけ別のヘッダー・
                  下線の記入欄・均等な段落間隔
  回覧.docx       蛍光ペン・取り消し線・上付き下付き・小さい大文字・
                  リンク・コメント・タブ位置・段落の罫線は使わない(様式は表で)

同じ5枚を officework エンジンで作るのは test/write_docx_officework.py。
突き合わせは目で見る(発注者 2026-08-26)。
"""
import sys
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT

OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).resolve().parent / "out"
OUT.mkdir(parents=True, exist_ok=True)

SKIPPED = []


def a4(doc, landscape=False):
    s = doc.sections[0]
    if landscape:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Mm(297), Mm(210)
    else:
        s.page_width, s.page_height = Mm(210), Mm(297)
    s.top_margin = s.bottom_margin = Mm(20)
    s.left_margin = s.right_margin = Mm(20)


# ======================================================================
# 1. 開催通知
# ======================================================================
def notice():
    d = Document()
    a4(d)
    p = d.add_paragraph("2026年8月26日")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = d.add_paragraph("会員各位")
    p.paragraph_format.space_after = Pt(12)
    p = d.add_paragraph("aiseed 事務局")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    h = d.add_heading("定例会 開催のお知らせ", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = d.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(10)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run("下記のとおり定例会を開きます。")
    r = p.add_run("出欠のご返事は 9月5日(金)まで")
    r.bold = True
    r.underline = True
    p.add_run("にお願いします。")

    p = d.add_paragraph("記")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    for label, body in [("日時", "9月12日(土) 14:00〜16:00"),
                        ("場所", "工業技術センター 2階 会議室"),
                        ("議題", "新しい道具の紹介と実演")]:
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Mm(20)
        p.paragraph_format.tab_stops.add_tab_stop(Mm(45), WD_TAB_ALIGNMENT.LEFT)
        r = p.add_run(label)
        r.bold = True
        p.add_run("\t" + body)

    p = d.add_paragraph("以上")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = d.add_paragraph("問い合わせ: 事務局(内線 123)")
    r = p.runs[0]
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    d.core_properties.title = "定例会の開催通知"
    d.core_properties.author = "write_docx_pydocx"
    d.save(OUT / "開催通知.docx")


# ======================================================================
# 2. 議事録
# ======================================================================
def minutes():
    d = Document()
    a4(d)
    h = d.add_heading("打ち合わせ議事録", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t = d.add_table(rows=4, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0, 0).text = "件名"
    a = t.cell(0, 1).merge(t.cell(0, 3))
    a.text = "玄関ドアのカタログ改善"
    t.cell(1, 0).text = "日時"
    t.cell(1, 1).text = "2026-08-26 10:00"
    t.cell(1, 2).text = "場所"
    t.cell(1, 3).text = "第2会議室"
    t.cell(2, 0).text = "出席"
    b = t.cell(2, 1).merge(t.cell(2, 3))
    b.text = "営業部2名・設計1名・事務局1名"
    t.cell(3, 0).text = "記録"
    t.cell(3, 1).text = "事務局"
    for row in t.rows:
        row.height = Mm(8)
        for c in row.cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for c in t.columns[0].cells:
        c.width = Mm(22)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    d.add_paragraph("決まったこと", style="Heading 2")
    for item in ["型と色を選ぶと合成図が出る形にする",
                 "寸法は物件ごとに mm で控える",
                 "次回までに見本の様式を3枚つくる"]:
        d.add_paragraph(item, style="List Bullet")
    d.add_paragraph("宿題", style="Heading 2")
    for item in ["様式の下書き(営業部)", "防火認定の一覧(設計)", "日程の調整(事務局)"]:
        d.add_paragraph(item, style="List Number")
    d.core_properties.title = "打ち合わせ議事録"
    d.save(OUT / "議事録.docx")


# ======================================================================
# 3. 操作手順書
# ======================================================================
def manual():
    d = Document()
    a4(d)
    # 等幅の文字スタイルを1つ作る
    st = d.styles.add_style("コマンドの字", WD_STYLE_TYPE.CHARACTER)
    st.font.name = "Courier New"
    st.font.size = Pt(10)

    d.add_heading("月次の締めの手順", level=1)
    d.add_paragraph("毎月1日の朝にやります。10分で終わります。")
    steps = [("台帳を開く", "先月のフォルダの 売上台帳.xlsx を開きます。"),
             ("マクロを押す", None),
             ("印刷する", "できた報告を A4 で1枚印刷し、回覧に載せます。")]
    for title, body in steps:
        d.add_paragraph(title, style="List Number")
        if body:
            p = d.add_paragraph(body)
        else:
            p = d.add_paragraph("マクロの一覧から ")
            p.add_run("月次の締め", style="コマンドの字")
            p.add_run(" を選んで押します。")
        p.paragraph_format.left_indent = Mm(10)

    logo = Path(__file__).resolve().parent.parent / "packaging/icons/hicolor/128x128/officework.png"
    if logo.exists():
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Mm(20))

    d.add_page_break()
    d.add_heading("別紙: 押すボタンの場所", level=2)
    p = d.add_paragraph("この行は2ページ目の頭にあります(改ページの確かめ)。")
    p.paragraph_format.keep_with_next = True
    d.core_properties.title = "月次の締めの手順"
    d.save(OUT / "操作手順書.docx")


# ======================================================================
# 4. 送付状
# ======================================================================
def cover_letter():
    d = Document()
    a4(d)
    s = d.sections[0]
    s.different_first_page_header_footer = True
    s.first_page_header.paragraphs[0].text = "aiseed — 送付状"
    s.header.paragraphs[0].text = "aiseed"
    f = s.footer.paragraphs[0]
    f.text = "この便りに心当たりが無いときは事務局へ"
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = d.add_paragraph("見積書 送付のご案内")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True
    p = d.add_paragraph("いつもお世話になっております。次の書類をお送りします。")
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)

    for name, n in [("御見積書", 1), ("カタログ(玄関ドア)", 1), ("返信用封筒", 1)]:
        p = d.add_paragraph(style="List Bullet")
        p.add_run(f"{name} … {n} 部")

    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.add_run("お受け取りの確認欄: ")
    r = p.add_run("                    ")
    r.underline = True
    p.add_run("(サイン)")
    d.core_properties.title = "送付状"
    d.save(OUT / "送付状.docx")


# ======================================================================
# 5. 回覧
# ======================================================================
def circular():
    d = Document()
    a4(d)
    h = d.add_heading("回覧 — 夏季の節電のお願い", level=1)

    p = d.add_paragraph()
    r = p.add_run("空調は 28℃ 設定")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run(" にご協力ください。")

    p = d.add_paragraph()
    p.add_run("昨年の指針(")
    r = p.add_run("27℃")
    r.font.strike = True
    p.add_run(" → 28℃)に改めています。")

    p = d.add_paragraph()
    p.add_run("面積あたりの目安は 15W/m")
    r = p.add_run("2")
    r.font.superscript = True
    p.add_run("、湿度は H")
    r = p.add_run("2")
    r.font.subscript = True
    p.add_run("O の量で変わります。")

    p = d.add_paragraph()
    r = p.add_run("Small Caps の見出し")
    r.font.small_caps = True

    p = d.add_paragraph("詳しい資料: ")
    if hasattr(p, "add_external_hyperlink"):
        p.add_external_hyperlink("https://github.com/aiseed-dev/officework", "officework の置き場")
    else:
        r = p.add_run("https://github.com/aiseed-dev/officework")
        r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        r.underline = True
        SKIPPED.append("リンクの挿入(この版の python-docx に口が無いので、字と色で代用)")

    p = d.add_paragraph("確認したら名前の右に日付を書いてください。")
    if hasattr(d, "add_comment"):
        d.add_comment(runs=p.runs, text="今週中にお願いします", author="事務局", initials="J")
    else:
        SKIPPED.append("コメント(この版の python-docx に口が無い)")

    t = d.add_table(rows=2, cols=5)
    t.style = "Table Grid"
    for j, name in enumerate(["佐藤", "鈴木", "高橋", "田中", "伊藤"]):
        t.cell(0, j).text = name
    d.core_properties.title = "回覧"
    d.save(OUT / "回覧.docx")


if __name__ == "__main__":
    notice()
    minutes()
    manual()
    cover_letter()
    circular()
    print("5枚書けた →", OUT)
    for s in SKIPPED:
        print("  口が無くて代用した:", s)
