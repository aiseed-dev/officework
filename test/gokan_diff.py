#!/usr/bin/env python3
"""**本家との差分の検査 — 本家の公開の口を全部歩く。**

同じ実物を本家(openpyxl / python-docx)と officework に読ませ、
**本家のクラスが公開しているプロパティを機械的に全部**突き合わせる
(2026-08-31 発注者「python-docx と openpyxl の全部を網羅したテストに」)。

見る項目を人が選ぶと、選ばなかった穴は永遠に見えない。ここは本家の
クラス定義から公開プロパティの一覧を取り出して回すので、本家に口が
増えれば検査も勝手に広がる。うちに口が無ければ「口が無い」、答えが
違えば「答えが違う」として、**口の名前ごと**に数える。

    .venv/bin/python test/gokan_diff.py ~/dev/test/zei/*.xlsx ~/dev/test/cao/*.docx

さらに**書きの往復**: officework で開いてそのまま保存した物を本家に
読ませ、本家が元から読んだ答えと同じ歩き方で比べる。触っていないのに
変わる物(原本を変える壊れ方)が出る。
"""

import inspect
import sys
import tempfile
from pathlib import Path

MAX_SHOW = 3   # 口ごとに最初の3例だけ見せる(数は全部数える)
MAX_ROWS = 400
MAX_COLS = 60
DEPTH = 3      # 値のオブジェクト(Font など)を掘る深さ


class Diff:
    def __init__(self, name):
        self.name = name
        self.kinds = {}

    def add(self, member, msg):
        self.kinds.setdefault(member, []).append(msg)

    def report(self):
        total = sum(len(v) for v in self.kinds.values())
        if not total:
            print(f"  {self.name}: ずれなし")
            return 0
        print(f"  {self.name}: ずれ {total} 件 / 口 {len(self.kinds)} 種")
        for member, msgs in sorted(self.kinds.items(), key=lambda kv: -len(kv[1])):
            print(f"    {member}: {len(msgs)} 件")
            for m in msgs[:MAX_SHOW]:
                print(f"      {m}")
            if len(msgs) > MAX_SHOW:
                print(f"      … 残り {len(msgs) - MAX_SHOW} 件")
        return total


# ---- 比べ方 ---------------------------------------------------------------

def props_of(cls) -> list:
    """本家のクラスの公開プロパティの名前。ここが網羅の正本 —
    本家の定義から取るので、人が項目を選ばない。"""
    out = []
    for n in dir(cls):
        if n.startswith("_"):
            continue
        try:
            a = inspect.getattr_static(cls, n)
        except AttributeError:
            continue
        if isinstance(a, (property, inspect.getattr_static(type("x", (), {"p": property(lambda s: 0)}), "p").__class__)):
            out.append(n)
    return sorted(set(out))


def norm(v, depth=DEPTH):
    """答えを比べられる形に。数の 100 と 100.0 は同じ、EMU は int、
    列挙は名前、値のオブジェクトは公開の場だけの辞書に。"""
    # **字を継いだ入れ物は、中身で比べます。** うちの `p.style` や
    # `run.font` は「字としても使える入れ物」なので、字として比べると
    # 本家の持ち物と突き合わせられません(2026-09-01)
    if isinstance(v, str) and type(v) is not str and (
            hasattr(v, "name") or hasattr(v, "rgb")):
        pass
    elif v is None or isinstance(v, (bool, str, bytes)):
        return v
    # 列挙(WD_ALIGN_PARAGRAPH.RIGHT など)は名前の小文字で —
    # うちは文字列で返す設計なので、表し方を揃えてから比べる
    if hasattr(v, "name") and hasattr(v, "__int__") and not isinstance(v, bool):
        try:
            return str(v.name).lower()
        except Exception:
            pass
    if isinstance(v, float):
        return int(v) if v == int(v) and abs(v) < 1e15 else round(v, 6)
    if isinstance(v, int):
        return int(v)
    if isinstance(v, (list, tuple)):
        return [norm(x, depth - 1) for x in v] if depth > 0 else f"<列 {len(v)}>"
    # **深さ切れの印にクラスの名前を使いません**(2026-09-01)。本家と
    # こちらで名前が同じになるはずがなく、中身を見ていないのに落ちます
    if depth <= 0:
        return "<入れ物>"
    # 値のオブジェクト(Font・Alignment・Color…)は公開の場を辞書に
    d = {}
    for n in dir(v):
        if n.startswith("_") or n in (
            "parent", "idx", "tagname", "namespace",
            # 内部の XML(lxml)は表し方の雑音にしかならない
            "element", "xml", "attrib", "nsmap", "prefix", "sourceline", "tag",
        ):
            continue
        try:
            a = getattr(v, n)
        except Exception:
            continue
        if callable(a):
            continue
        d[n] = norm(a, depth - 1)
    return d or "<入れ物>"


def cmp_prop(d: Diff, where: str, member: str, ref, mine):
    """本家の答えとうちの答えを1つ比べる。"""
    try:
        rv = getattr(ref, member)
    except Exception:
        return  # 本家自身が答えられない口は比べない
    if callable(rv):
        return
    try:
        mv = getattr(mine, member)
    except AttributeError:
        d.add(f"口が無い: {member}", where)
        return
    except Exception as e:
        d.add(f"口が壊れる: {member}", f"{where}: {type(e).__name__} {e}")
        return
    if callable(mv):
        mv_desc = "<関数>"
        d.add(f"形が違う: {member}", f"{where}: 本家は値 / うちは{mv_desc}")
        return
    a, b = norm(rv), norm(mv)
    if a != b:
        sa, sb = repr(a)[:60], repr(b)[:60]
        d.add(f"答えが違う: {member}", f"{where}: 本家 {sa} / うち {sb}")


def walk_pair(d: Diff, where: str, ref, mine, skip=()):
    """本家のクラスの公開プロパティを全部歩いて比べる。"""
    for member in props_of(type(ref)):
        if member in skip:
            continue
        cmp_prop(d, where, member, ref, mine)


# ---- xlsx -----------------------------------------------------------------

# 比べても意味の無い口(こちらの内部や、実物でなく環境に依る物)
XLSX_SKIP_WB = {"loaded_theme", "vba_archive", "path", "excel_base_date",
                "epoch", "template", "data_only", "read_only", "write_only",
                "iso_dates", "rels", "calculation", "views", "security",
                "shared_strings", "style_names"}
XLSX_SKIP_WS = {"parent", "views", "HeaderFooter", "legacy_drawing",
                "orientation", "path", "plot"}
# **意図した違い**(python-manual の「意図した違い」に書いたもの)。
# 理由を書かずにここへ足さないこと。
#
# style_id / has_style — openpyxl が自分の中に持つ書式表の番号です。
#   同じ見た目でも並べ方で番号が変わるので、突き合わせても意味がありません。
#   こちらは升ごとに書式そのものを持ちます(`cell.font` などで読めます)。
#   openpyxl は結合の2升目以降を「書式なし」と答えますが、ファイルには
#   書式の番号が書いてあります。こちらはファイルのとおりに答えます。
XLSX_SKIP_CELL = {"parent", "encoding", "base_date", "style_id", "has_style"}


def xlsx_diff(path: Path) -> int:
    import openpyxl
    from officework import sheet

    d = Diff(path.name)
    wb = openpyxl.load_workbook(path)
    b = sheet.Book.open(str(path))

    walk_pair(d, "Workbook", wb, b, skip=XLSX_SKIP_WB)
    for name in wb.sheetnames:
        if name not in b.sheet_names:
            d.add("シートが無い", name)
            continue
        ws, s = wb[name], b[name]
        walk_pair(d, f"[{name}]", ws, s, skip=XLSX_SKIP_WS)
        rows = min(ws.max_row, MAX_ROWS)
        cols = min(ws.max_column, MAX_COLS)
        for r in range(1, rows + 1):
            for c in range(1, cols + 1):
                oc = ws.cell(r, c)
                try:
                    mc = s.cell(r, c)
                except Exception as e:
                    d.add("cell() が壊れる", f"{name}!{oc.coordinate}: {e}")
                    continue
                walk_pair(d, f"{name}!{oc.coordinate}", oc, mc, skip=XLSX_SKIP_CELL)
    return d.report()


def xlsx_roundtrip(path: Path) -> int:
    """officework で開いてそのまま保存 → 本家どうしで元と比べる。"""
    import openpyxl
    from officework import sheet

    d = Diff(path.name + "(往復)")
    out = Path(tempfile.mkstemp(suffix=".xlsx")[1])
    b = sheet.Book.open(str(path))
    b.save(str(out))
    a = openpyxl.load_workbook(path)
    z = openpyxl.load_workbook(out)
    for name in a.sheetnames:
        if name not in z.sheetnames:
            d.add("シートが消える", name)
            continue
        wa, wz = a[name], z[name]
        rows = min(wa.max_row, MAX_ROWS)
        cols = min(wa.max_column, MAX_COLS)
        for r in range(1, rows + 1):
            for c in range(1, cols + 1):
                ca, cz = wa.cell(r, c), wz.cell(r, c)
                walk_pair(d, f"{name}!{ca.coordinate}", ca, cz, skip=XLSX_SKIP_CELL)
    out.unlink()
    return d.report()


# ---- docx -----------------------------------------------------------------

DOCX_SKIP_DOC = {"element", "part", "settings", "styles"}
DOCX_SKIP_PARA = {"part"}
DOCX_SKIP_RUN = {"part", "element"}
DOCX_SKIP_SEC = set()


def docx_diff(path: Path) -> int:
    import docx
    from officework import doc as odoc

    d = Diff(path.name)
    a = docx.Document(str(path))
    m = odoc.Doc.open(str(path))

    apara = list(a.paragraphs)
    if len(apara) != len(m):
        d.add("答えが違う: len(paragraphs)", f"本家 {len(apara)} / うち {len(m)}")
    for i in range(min(len(apara), len(m))):
        pa, pm = apara[i], m[i]
        walk_pair(d, f"段落{i}", pa, pm, skip=DOCX_SKIP_PARA)
        ra, rm = list(pa.runs), list(pm.runs)
        if len(ra) != len(rm):
            d.add("答えが違う: len(runs)", f"段落{i}: 本家 {len(ra)} / うち {len(rm)}")
        for j in range(min(len(ra), len(rm))):
            walk_pair(d, f"段落{i}.run{j}", ra[j], rm[j], skip=DOCX_SKIP_RUN)
    ta, tm = list(a.tables), list(m.tables)
    if len(ta) != len(tm):
        d.add("答えが違う: len(tables)", f"本家 {len(ta)} / うち {len(tm)}")
    for ti in range(min(len(ta), len(tm))):
        rows_a = ta[ti].rows
        for r in range(min(len(rows_a), len(tm[ti]))):
            ca = [c.text for c in rows_a[r].cells]
            try:
                cm = [tm[ti][r][c].text for c in range(len(tm[ti][r]))]
            except Exception as e:
                d.add("口が壊れる: table[r][c]", f"表{ti}行{r}: {e}")
                continue
            if ca != cm:
                d.add("答えが違う: table cell text",
                      f"表{ti}行{r}: 本家 {ca[:2]!r} / うち {cm[:2]!r}")
    sa, sm = list(a.sections), list(m.sections)
    if len(sa) != len(sm):
        d.add("答えが違う: len(sections)", f"本家 {len(sa)} / うち {len(sm)}")
    for i in range(min(len(sa), len(sm))):
        walk_pair(d, f"節{i}", sa[i], sm[i], skip=DOCX_SKIP_SEC)
    return d.report()


def docx_roundtrip(path: Path) -> int:
    import docx
    from officework import doc as odoc

    d = Diff(path.name + "(往復)")
    out = Path(tempfile.mkstemp(suffix=".docx")[1])
    m = odoc.Doc.open(str(path))
    m.save(str(out))
    a = docx.Document(str(path))
    z = docx.Document(str(out))
    pa, pz = list(a.paragraphs), list(z.paragraphs)
    if len(pa) != len(pz):
        d.add("段落の数が変わる", f"{len(pa)} → {len(pz)}")
    for i in range(min(len(pa), len(pz))):
        walk_pair(d, f"段落{i}", pa[i], pz[i], skip=DOCX_SKIP_PARA)
    if len(a.tables) != len(z.tables):
        d.add("表の数が変わる", f"{len(a.tables)} → {len(z.tables)}")
    for i in range(min(len(a.sections), len(z.sections))):
        walk_pair(d, f"節{i}", a.sections[i], z.sections[i], skip=DOCX_SKIP_SEC)
    out.unlink()
    return d.report()


def main():
    paths = [Path(p).expanduser() for p in sys.argv[1:]]
    if not paths:
        sys.exit("使い方: gokan_diff.py <実物.xlsx/.docx>…")
    total = 0
    for p in paths:
        print(f"== {p}")
        if p.suffix.lower() == ".xlsx":
            total += xlsx_diff(p)
            total += xlsx_roundtrip(p)
        elif p.suffix.lower() == ".docx":
            total += docx_diff(p)
            total += docx_roundtrip(p)
    print(f"\n合計のずれ: {total} 件")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
