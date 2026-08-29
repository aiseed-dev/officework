#!/usr/bin/env python3
"""文書の図形の基本のテスト。

    .venv/bin/python test/basic_docshape.py

python-docx には無い口です(本家は画像しか置けません)。docx にも紙にも
出ること、Word や LibreOffice が図形として開ける形で書けることを見ます。
"""
import tempfile
import zipfile
from pathlib import Path

from officework import doc

tmp = Path(tempfile.mkdtemp(prefix="basic-docshape-"))
BAD = []


def check(label, got, want):
    if got != want:
        BAD.append(label)
        print(f"× {label}: {want!r} のはずが {got!r}")
    else:
        print(f"  {label}: OK")


def moto():
    d = doc.Doc()
    d.add_heading("図形の入った文書", level=1)
    d.add_paragraph("本文です。")
    return d


# 1. 6種類とも置ける
d = moto()
for i, kind in enumerate(
    ("rect", "roundRect", "ellipse", "rightArrow", "diamond", "line")
):
    d.add_shape(kind, 20.0 + i * 5.0, 80.0, 40.0, 25.0, fill="DDE7F0", line="2E5A87")
check("6種類が置ける", d.shapes, 6)

# 2. 知らない形は断る。**できないことをできるように見せない**
try:
    moto().add_shape("smiley", 10.0, 10.0, 20.0, 20.0)
    check("知らない形を断る", "受けてしまった", "断る")
except ValueError:
    check("知らない形を断る", "断る", "断る")

# 3. 大きさが 0 なら断る
try:
    moto().add_shape("rect", 10.0, 10.0, 0.0, 20.0)
    check("大きさ 0 を断る", "受けてしまった", "断る")
except ValueError:
    check("大きさ 0 を断る", "断る", "断る")

# 4. docx に**図形として**書かれる(絵に落とさない)
d = moto()
d.add_shape("rect", 25.0, 80.0, 40.0, 25.0, fill="DDE7F0", line="2E5A87", text="四角")
d.add_shape("roundRect", 80.0, 80.0, 40.0, 25.0, fill="F5E6D3", shadow=True)
d.add_shape("ellipse", 135.0, 80.0, 40.0, 25.0, fill="C0504D", opacity=0.5)
p = tmp / "shape.docx"
d.save(str(p))
x = zipfile.ZipFile(p).read("word/document.xml").decode("utf-8")
check("docx の頭", p.read_bytes()[:2], b"PK")
check("図形が3つ", x.count("<wps:wsp>"), 3)
check("形の名前が入る", 'prst="roundRect"' in x and 'prst="ellipse"' in x, True)
check("図形の中の文字", "四角" in x, True)
check("影", "outerShdw" in x, True)
check("不透明度", "a:alpha" in x, True)

# 5. python-docx で開ける(壊れた docx を書いていない)
import docx as pdocx

check("python-docx で開ける", len(pdocx.Document(str(p)).paragraphs) >= 2, True)

# 6. 紙にも出る
q = tmp / "shape.pdf"
d.save(str(q))
check("PDF の頭", q.read_bytes()[:5], b"%PDF-")
check("PDF に中身がある", q.stat().st_size > 1000, True)

if BAD:
    raise SystemExit(f"合わない物 {len(BAD)} 件: " + "、".join(BAD))
print("基本のテスト: 全部合った")
