#!/usr/bin/env python3
"""**Microsoft 365 で教師データを作るための、1枚1機能の見本。**

来月の M365 Personal の契約(2026-09-01 発注者)に合わせた支度。
Mac の Word / Excel でこの見本を開いて PDF に書き出せば、機能ごとの
「本物はこう組む」が手に入り、酒税・告知書でやった見比べを機能単位で
機械的に回せるようになる。

**見本は本家(openpyxl / python-docx)で生成する。** officework で
作るとうちの癖が教師に混ざるため。

    .venv/bin/python test/gen_ms_teacher.py ~/dev/test/ms

Mac での手順(出力先の README にも書く):

1. 各ファイルを Word / Excel で開く
2. そのまま PDF に書き出す(ファイル > 名前を付けて保存 > PDF)。
   名前は <元の名前>_ms.pdf、置き場は同じフォルダ
3. officework 側は同じ物を <元の名前>_ow.pdf に(道具は今後)
4. ずれたら、その機能の名前がそのまま直しの題名になる
"""

import sys
from pathlib import Path

DEFAULT_OUT = Path.home() / "dev/test/ms"


# ---------------- xlsx(1枚1機能) ----------------

def gen_xlsx(out: Path):
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    def book():
        wb = openpyxl.Workbook()
        return wb, wb.active

    def save(wb, name):
        p = out / name
        wb.save(p)
        print(" ", p.name)

    # 1. 罫線の線種13種
    wb, ws = book()
    styles = ["hair", "dotted", "dashDotDot", "dashDot", "dashed", "thin",
              "mediumDashDotDot", "mediumDashDot", "mediumDashed", "medium",
              "thick", "double", "slantDashDot"]
    for i, st in enumerate(styles, start=2):
        ws.cell(i, 2, st)
        ws.cell(i, 4).border = Border(bottom=Side(style=st))
        ws.column_dimensions["D"].width = 20
    save(wb, "x01_罫線の線種.xlsx")

    # 2. 結合セル(横・縦・面)と外枠
    wb, ws = book()
    ws.merge_cells("B2:E2")
    ws["B2"] = "横に4つ結合"
    ws.merge_cells("B4:B7")
    ws["B4"] = "縦に4つ結合"
    ws.merge_cells("D4:F7")
    ws["D4"] = "面で結合"
    thin = Side(style="thin")
    for rng in ("B2:E2", "B4:B7", "D4:F7"):
        for row in ws[rng]:
            for c in row:
                c.border = Border(top=thin, bottom=thin, left=thin, right=thin)
    save(wb, "x02_結合セル.xlsx")

    # 3. 表示形式(¥・%・日付・負数の△・小数)
    import datetime
    wb, ws = book()
    rows = [
        (640200, "¥#,##0"),
        (0.123, "0.0%"),
        (datetime.date(2026, 9, 1), "yyyy\"年\"m\"月\"d\"日\""),
        (-5148, '#,##0;"△ "#,##0'),
        (3.14159, "0.00"),
        (1234567.89, "#,##0.00"),
    ]
    for i, (v, fmt) in enumerate(rows, start=2):
        ws.cell(i, 2, v).number_format = fmt
        ws.cell(i, 4, fmt)
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["D"].width = 24
    save(wb, "x03_表示形式.xlsx")

    # 4. 縮小印刷(scale=69。酒税の表で 22 ページになった機能)
    wb, ws = book()
    for r in range(1, 61):
        for c in range(1, 16):
            ws.cell(r, c, f"{get_column_letter(c)}{r}")
    ws.page_setup.scale = 69
    save(wb, "x04_縮小印刷69.xlsx")

    # 5. 紙1枚に収める(fitToPage)
    wb, ws = book()
    for r in range(1, 81):
        for c in range(1, 12):
            ws.cell(r, c, r * c)
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    save(wb, "x05_紙1枚に収める.xlsx")

    # 6. 行の高さ・列の幅の指定
    wb, ws = book()
    ws["B2"] = "高さ30の行"
    ws.row_dimensions[2].height = 30
    ws["B4"] = "高さ8の行"
    ws.row_dimensions[4].height = 8
    ws["D2"] = "幅30の列"
    ws.column_dimensions["D"].width = 30
    save(wb, "x06_行高と列幅.xlsx")

    # 7. 書体の混在(明朝・ゴシック・Century・大きさ)
    wb, ws = book()
    cells = [
        ("ＭＳ 明朝の11pt", Font(name="ＭＳ 明朝", size=11)),
        ("ＭＳ ゴシックの11pt", Font(name="ＭＳ ゴシック", size=11)),
        ("Century 11pt", Font(name="Century", size=11)),
        ("游明朝の14pt", Font(name="游明朝", size=14)),
        ("太字の明朝", Font(name="ＭＳ 明朝", bold=True)),
    ]
    for i, (v, f) in enumerate(cells, start=2):
        ws.cell(i, 2, v).font = f
    ws.column_dimensions["B"].width = 28
    save(wb, "x07_書体の混在.xlsx")

    # 8. 折り返しと縮小して全体を表示
    wb, ws = book()
    long = "折り返して全体を表示する長い文章がこのセルに入っています"
    ws["B2"] = long
    ws["B2"].alignment = Alignment(wrap_text=True)
    ws["B4"] = long
    ws["B4"].alignment = Alignment(shrink_to_fit=True)
    ws["B6"] = "均等割付の字"
    ws["B6"].alignment = Alignment(horizontal="distributed")
    ws.column_dimensions["B"].width = 16
    save(wb, "x08_折り返しと均等割付.xlsx")

    # 9. 塗りと文字色
    wb, ws = book()
    ws["B2"] = "黄色の地"
    ws["B2"].fill = PatternFill("solid", fgColor="FFFF00")
    ws["B4"] = "赤い字"
    ws["B4"].font = Font(color="FF0000")
    ws["B6"] = "青地に白"
    ws["B6"].fill = PatternFill("solid", fgColor="1F4E79")
    ws["B6"].font = Font(color="FFFFFF")
    save(wb, "x09_塗りと文字色.xlsx")

    # 10. 印刷範囲とタイトル行(2ページ目にも見出し)
    wb, ws = book()
    ws["A1"] = "見出しの行(全ページに出る)"
    for r in range(2, 91):
        ws.cell(r, 1, f"明細 {r - 1}")
    ws.print_area = "A1:C90"
    ws.print_title_rows = "1:1"
    save(wb, "x10_印刷範囲とタイトル行.xlsx")


# ---------------- docx(1枚1機能) ----------------

def gen_docx(out: Path):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt, RGBColor

    def save(d, name):
        p = out / name
        d.save(p)
        print(" ", p.name)

    # 1. 行送り(exact / atLeast / 倍率。beta.7 の題材)
    d = docx.Document()
    p = d.add_paragraph("行送りを exact 20pt にした段落です。" * 3)
    p.paragraph_format.line_spacing = Pt(20)
    p = d.add_paragraph("行送りを 1.5 倍にした段落です。" * 3)
    p.paragraph_format.line_spacing = 1.5
    p = d.add_paragraph("行送りを指定しない段落です。" * 3)
    save(d, "d01_行送り.docx")

    # 2. 字下げ(インデント指定と、行頭の全角スペース)
    d = docx.Document()
    p = d.add_paragraph("インデント指定で2字ぶん下げた段落です。" * 2)
    p.paragraph_format.first_line_indent = Mm(7.4)
    d.add_paragraph("　行頭に全角スペースを打った段落です(日本の書類の普通)。" * 2)
    d.add_paragraph("字下げの無い段落です。" * 2)
    save(d, "d02_字下げ.docx")

    # 3. 揃え(左・中央・右・両端)
    d = docx.Document()
    for text, al in [("左揃え", WD_ALIGN_PARAGRAPH.LEFT),
                     ("中央揃え", WD_ALIGN_PARAGRAPH.CENTER),
                     ("右揃え", WD_ALIGN_PARAGRAPH.RIGHT),
                     ("両端揃えの長い文章。" * 6, WD_ALIGN_PARAGRAPH.JUSTIFY)]:
        p = d.add_paragraph(text)
        p.alignment = al
    save(d, "d03_揃え.docx")

    # 4. 文字の書式(太字・斜体・下線・色・大きさ・run の混在)
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("普通の字、")
    p.add_run("太字、").bold = True
    p.add_run("斜体、").italic = True
    p.add_run("下線、").underline = True
    r = p.add_run("赤い字、")
    r.font.color.rgb = RGBColor(0xFF, 0, 0)
    r = p.add_run("大きい字。")
    r.font.size = Pt(16)
    save(d, "d04_文字の書式.docx")

    # 5. 見出しのスタイル(定義の色と大きさが効くか。beta.7 の題材)
    d = docx.Document()
    d.add_heading("見出し1の字", level=1)
    d.add_paragraph("本文の字。" * 5)
    d.add_heading("見出し2の字", level=2)
    d.add_paragraph("本文の字。" * 5)
    save(d, "d05_見出しスタイル.docx")

    # 6. 表(結合・罫線・列幅)
    d = docx.Document()
    t = d.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    t.cell(0, 0).merge(t.cell(0, 2))
    t.cell(0, 0).text = "3列を結合した見出し"
    for r in range(1, 3):
        for c in range(3):
            t.cell(r, c).text = f"{r}-{c}"
    save(d, "d06_表.docx")

    # 7. 画像(字と同じ行に。beta.7 の題材)— 小さな PNG をその場で作る
    import struct
    import zlib as z

    def png_bytes(w, h, rgb):
        raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
        def chunk(tag, data):
            c = struct.pack(">I", len(data)) + tag + data
            return c + struct.pack(">I", z.crc32(tag + data) & 0xFFFFFFFF)
        return (b"\x89PNG\r\n\x1a\n"
                + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
                + chunk(b"IDAT", z.compress(raw))
                + chunk(b"IEND", b""))

    img = out / "_四角.png"
    img.write_bytes(png_bytes(60, 40, (0x2E, 0x5A, 0x87)))
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("字のとなりに ")
    p.add_run().add_picture(str(img), width=Mm(15))
    p.add_run(" 絵が入る行です。")
    save(d, "d07_字と画像.docx")

    # 8. 節(途中で用紙が変わる)
    from docx.enum.section import WD_SECTION
    d = docx.Document()
    d.add_paragraph("1節目は A4 縦です。")
    s = d.add_section(WD_SECTION.NEW_PAGE)
    s.page_width, s.page_height = Mm(297), Mm(210)  # A4 横
    d.add_paragraph("2節目は A4 横です。")
    save(d, "d08_節で用紙が変わる.docx")

    # 9. ヘッダーとフッター(ページ番号は Word で振る)
    d = docx.Document()
    d.sections[0].header.paragraphs[0].text = "ヘッダーの字"
    d.sections[0].footer.paragraphs[0].text = "フッターの字"
    for i in range(1, 4):
        d.add_paragraph(f"{i} ページ目の本文。" * 40)
    save(d, "d09_ヘッダーフッター.docx")

    # 10. 改ページ
    d = docx.Document()
    d.add_paragraph("1ページ目。")
    d.add_page_break()
    d.add_paragraph("2ページ目(改ページで来た)。")
    save(d, "d10_改ページ.docx")


def main():
    out = Path(sys.argv[1]).expanduser() if len(sys.argv) > 1 else DEFAULT_OUT
    out.mkdir(parents=True, exist_ok=True)
    print(f"出力: {out}")
    gen_xlsx(out)
    gen_docx(out)
    (out / "README.md").write_text(
        "# Microsoft 365 で教師データを作る手順\n\n"
        "1枚1機能の見本(x01〜x10 が Excel、d01〜d10 が Word)。\n"
        "本家の openpyxl / python-docx で生成してある(生成し直しは\n"
        "officework の test/gen_ms_teacher.py)。\n\n"
        "1. 各ファイルを Word / Excel で開く\n"
        "2. そのまま PDF に書き出す(名前は <元の名前>_ms.pdf、同じフォルダへ)\n"
        "3. officework でも同じ物を <元の名前>_ow.pdf に\n"
        "4. 2つの PDF を見比べる。ずれたら、ファイル名の機能がそのまま直しの題名\n",
        encoding="utf-8",
    )
    print("README.md も書いた")


if __name__ == "__main__":
    main()
