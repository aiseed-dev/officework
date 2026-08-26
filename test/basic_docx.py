#!/usr/bin/env python3
"""文書の基本のテスト — 毎日する操作だけ。python-docx の目でも読み返す。

    .venv/bin/python test/basic_docx.py
"""
import tempfile
from pathlib import Path

from officework import doc

tmp = Path(tempfile.mkdtemp(prefix="basic-docx-"))
BAD = []


def check(label, got, want):
    if got != want:
        BAD.append(label)
        print(f"× {label}: {want!r} のはずが {got!r}")
    else:
        print(f"  {label}: OK")


# 1. 段落と字の書式
d = doc.Doc()
p = d.add_paragraph("こんにちは。")
r = p.add_run("大事な字")
r.bold = True
p2 = d.add_paragraph("2つ目の段落です。")
check("段落の字", d.paragraphs[0].text, "こんにちは。大事な字")
check("太字", d.paragraphs[0].runs[-1].bold, True)
check("段落の数", len(d.paragraphs), 2)

# 2. 見出し
h = d.add_heading("章の題", level=1)
check("見出しの字", d.paragraphs[-1].text, "章の題")

# 3. 表 — 値の読み書き
t = d.add_table(rows=2, cols=2)
t.cell(0, 0).text = "品名"
t.cell(0, 1).text = "数"
t.cell(1, 0).text = "ペン"
t.cell(1, 1).text = "5"
check("表の値", d.tables[0].cell(1, 0).text, "ペン")

# 4. 保存して officework で開いて同じ
pa = str(tmp / "basic.docx")
d.save(pa)
d2 = doc.Doc.open(pa)
check("開き直した段落", d2.paragraphs[0].text, "こんにちは。大事な字")
check("開き直した太字", d2.paragraphs[0].runs[-1].bold, True)
check("開き直した見出しの字", d2.paragraphs[-1].text, "章の題")
check("開き直した表", d2.tables[0].cell(1, 1).text, "5")
check("読めなかった物", list(d2.unsupported), [])

# 5. python-docx の目で読み返して同じ
import docx as pydocx
w = pydocx.Document(pa)
check("本家の目(段落)", w.paragraphs[0].text, "こんにちは。大事な字")
check("本家の目(太字)", w.paragraphs[0].runs[-1].bold, True)
check("本家の目(見出しの様式)", w.paragraphs[-1].style.name, "Heading 1")
check("本家の目(表)", w.tables[0].cell(1, 0).text, "ペン")

# 6. 逆向き — python-docx が書いた物を officework で読む
q = pydocx.Document()
q.add_paragraph("本家が書いた段落。")
tq = q.add_table(rows=1, cols=2)
tq.cell(0, 0).text = "甲"
tq.cell(0, 1).text = "乙"
pb = str(tmp / "pydocx.docx")
q.save(pb)
d3 = doc.Doc.open(pb)
check("逆向き(段落)", d3.paragraphs[0].text, "本家が書いた段落。")
check("逆向き(表)", d3.tables[0].cell(0, 1).text, "乙")

# 7. 実物 — 開いて保存して、本家の目で数が合う
src = Path(__file__).resolve().parent.parent / "sample/報告書.docx"
if src.exists():
    d4 = doc.Doc.open(str(src))
    n_para = len(d4.paragraphs)
    pc = str(tmp / "roundtrip.docx")
    d4.save(pc)
    w2 = pydocx.Document(pc)
    check("実物の往復(段落の数)", len(w2.paragraphs), n_para)
else:
    print("  (sample/報告書.docx が無いので実物の往復は飛ばした)")

if BAD:
    raise SystemExit(f"合わない物 {len(BAD)} 件: " + "、".join(BAD))
print("基本のテスト: 全部合った")
